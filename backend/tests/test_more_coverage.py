import json

import pytest
from app import llm_client
from app.llm_config import _LLM_SETTINGS_PATH, save_llm_settings
from app.services import resume_gen as rg
from app.services.resume_gen import TEMPLATE_DIR
from app.routers import llm as llm_router


# ---------- llm_client: stream 解析 ----------

class _FakeStreamResp:
    def __init__(self, status=200, lines=(), body=b""):
        self.status_code = status
        self._lines = list(lines)
        self._body = body
    async def __aenter__(self):
        return self
    async def __aexit__(self, *a):
        return False
    async def aiter_lines(self):
        for ln in self._lines:
            yield ln
    async def aread(self):
        return self._body


class _FakeStreamCM:
    def __init__(self, resp):
        self._resp = resp
    async def __aenter__(self):
        return self._resp
    async def __aexit__(self, *a):
        return False


class _FakeStreamClient:
    def __init__(self, resp):
        self._resp = resp
    async def __aenter__(self):
        return self
    async def __aexit__(self, *a):
        return False
    def stream(self, method, url, json=None, headers=None):
        return _FakeStreamCM(self._resp)


@pytest.mark.asyncio
async def test_deepseek_stream_success(monkeypatch):
    import httpx
    lines = [
        'data: {"choices":[{"delta":{"content":"你"}}]}',
        'data: {"choices":[{"delta":{"content":"好"}}]}',
        'data: {"choices":[{"delta":{}}]}',
        'data: [DONE]',
    ]
    monkeypatch.setattr(httpx, "AsyncClient", lambda timeout=None: _FakeStreamClient(_FakeStreamResp(200, lines)))
    save_llm_settings({"provider": "deepseek", "deepseek_api_key": "sk-x"})
    cfg = llm_client.get_llm_settings()
    tokens = [t async for t in llm_client._deepseek_stream(cfg, "p", "sys", 0.7)]
    assert tokens == ["你", "好"]


@pytest.mark.asyncio
async def test_deepseek_stream_http_error(monkeypatch):
    import httpx
    monkeypatch.setattr(httpx, "AsyncClient", lambda timeout=None: _FakeStreamClient(_FakeStreamResp(500, [], b"boom")))
    save_llm_settings({"provider": "deepseek", "deepseek_api_key": "sk-x"})
    with pytest.raises(RuntimeError, match="500"):
        _ = [t async for t in llm_client._deepseek_stream(llm_client.get_llm_settings(), "p")]


def test_llm_client_helpers():
    msgs = llm_client._deepseek_messages("系统提示", "用户问题")
    assert msgs[0] == {"role": "system", "content": "系统提示"}
    assert msgs[1]["role"] == "user"
    assert llm_client._deepseek_messages("", "q") == [{"role": "user", "content": "q"}]
    h = llm_client._deepseek_headers("sk-1")
    assert h["Authorization"] == "Bearer sk-1"


# ---------- resume_gen: 表格 / 页眉占位 ----------

def _build_complex_template() -> str:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    doc.add_paragraph("{{姓名}}")
    tb = doc.add_table(rows=1, cols=2)
    tb.cell(0, 0).text = "{{电话}}"
    tb.cell(0, 1).text = "普通"
    header = doc.sections[0].header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.add_run("机密：{{求职意向}}")
    doc.add_paragraph("{{页眉标记}}")
    p = TEMPLATE_DIR / "complex.docx"
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(p))
    return str(p)


def test_complex_template_parse_and_merge():
    path = _build_complex_template()
    fields = rg.parse_template_fields(path)
    assert {"姓名", "电话", "求职意向", "页眉标记"} <= set(fields)

    out = TEMPLATE_DIR / "complex_out.docx"
    values = {"姓名": "张三", "电话": "123", "求职意向": "后端", "页眉标记": "示例"}
    rg.merge_docx(path, out, values)
    from docx import Document
    doc = Document(str(out))
    all_text = "\n".join([p.text for p in doc.paragraphs] + [t.text for t in [tb.cell(0, 0) for tb in doc.tables]])
    all_text += doc.sections[0].header.paragraphs[0].text if doc.sections[0].header.paragraphs else ""
    assert "张三" in all_text and "机密：后端" in all_text
    assert "{{" not in all_text


def test_merge_keeps_unknown_placeholder():
    path = _build_complex_template()
    out = TEMPLATE_DIR / "complex_unknown.docx"
    rg.merge_docx(path, out, {"姓名": "李四"})  # 其余字段缺失
    from docx import Document
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "李四" in text
    assert "{{页眉标记}}" in text  # 未知/缺失占位（正文里）保留


def test_save_to_desktop_disabled():
    # 测试环境 APP_RESUME_SAVE_DESKTOP=0 -> 不会写真实桌面
    p = TEMPLATE_DIR / "tmp_doc.docx"
    rg.build_example_template(p)
    assert rg.save_to_desktop(p, "简历-某人.docx") is None


def test_generate_docx_errors():
    with pytest.raises(ValueError):
        rg.generate_docx("missing-session")
    tpl = rg.save_template("gt.docx", _sample_bytes())
    s = rg.create_session(tpl["template_id"], None)
    # 删除模板文件 -> 触发“模板文件缺失”
    import shutil
    src = TEMPLATE_DIR / f"{tpl['template_id']}.docx"
    shutil.move(str(src), str(src) + ".bak")
    try:
        with pytest.raises(ValueError, match="模板文件缺失"):
            rg.generate_docx(s["session_id"])
    finally:
        shutil.move(str(src) + ".bak", str(src))


def _sample_bytes() -> bytes:
    from pathlib import Path
    p = TEMPLATE_DIR / "coversample.docx"
    rg.build_example_template(p)
    return Path(p).read_bytes()


# ---------- llm_config 归一化/异常 ----------

def test_config_corrupt_file_falls_back():
    _LLM_SETTINGS_PATH.parent.mkdir(parents=True, exist_ok=True)
    _LLM_SETTINGS_PATH.write_text("{broken", encoding="utf-8")
    cfg = llm_client.get_llm_settings()
    assert cfg.provider == "ollama"


def test_config_invalid_provider_normalized():
    _LLM_SETTINGS_PATH.parent.mkdir(parents=True, exist_ok=True)
    _LLM_SETTINGS_PATH.write_text(json.dumps({"provider": "openai", "deepseek_model": "", "deepseek_base_url": ""}),
                                  encoding="utf-8")
    cfg = llm_client.get_llm_settings()
    assert cfg.provider == "ollama"
    assert cfg.deepseek_model == "deepseek-v4-flash"
    assert cfg.deepseek_base_url == "https://api.deepseek.com"


# ---------- router 错误分支 ----------

def test_resume_gen_example_internal_error(client, enable_api, monkeypatch):
    def boom(*a, **k):
        raise RuntimeError("build failed")
    monkeypatch.setattr(rg, "build_example_template", boom)
    assert client.get("/api/v1/resume-gen/example-template").status_code == 500


def test_resume_gen_chat_and_generate_errors(client, enable_api):
    # 空消息 400
    sid = _make_session(client)
    assert client.post(f"/api/v1/resume-gen/sessions/{sid}/chat", json={"message": "  "}).status_code == 400
    # 下载但未生成 404
    assert client.get(f"/api/v1/resume-gen/sessions/{sid}/download").status_code == 404
    # 生成不存在会话 400
    assert client.post("/api/v1/resume-gen/sessions/nope/generate").status_code == 400
    # 上传非 docx 扩展名（PK 头但扩展名不对）
    r = client.post("/api/v1/resume-gen/upload-template",
                    files={"file": ("t.txt", b"PK\x03\x04xx", "application/octet-stream")})
    assert r.status_code == 400


def _make_session(client) -> str:
    ex = client.get("/api/v1/resume-gen/example-template").content
    r = client.post("/api/v1/resume-gen/upload-template",
                    files={"file": ("tmpl.docx", ex, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    sid = client.post("/api/v1/resume-gen/sessions", json={"template_id": r.json()["template_id"]}).json()["session_id"]
    return sid


def test_mbti_wrong_dim_counts(client, enable_api):
    answers = []
    # EI 给 6 题，其余维度凑 14 题 -> 总数 20 但每维不是 5
    answers += [{"dim": "EI", "pole": "E"}] * 6
    for dim, pole in [("SN", "S"), ("TF", "T"), ("JP", "J")]:
        answers += [{"dim": dim, "pole": pole}] * 5
    r = client.post("/api/v1/mbti/result", json={"answers": answers})
    assert r.status_code == 400


def test_llm_settings_save_500(client, monkeypatch):
    def boom(*a, **k):
        raise RuntimeError("disk full")
    monkeypatch.setattr(llm_router, "save_llm_settings", boom)
    r = client.put("/api/v1/llm/settings", json={"provider": "ollama"})
    assert r.status_code == 500

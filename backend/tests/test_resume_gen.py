import pytest
from app.services import resume_gen as rg
from app.services.resume_gen import TEMPLATE_DIR


# ---------- docx / service ----------

def _build_example() -> str:
    rg._ensure_dirs()
    p = TEMPLATE_DIR / "sample.docx"
    rg.build_example_template(p)
    return str(p)


def _sample_bytes() -> bytes:
    _build_example()
    return (TEMPLATE_DIR / "sample.docx").read_bytes()


def test_example_fields_and_merge():
    path = _build_example()
    fields = rg.parse_template_fields(path)
    assert len(fields) >= 10
    assert "姓名" in fields and "个人简介" in fields

    out = TEMPLATE_DIR / "out.docx"
    values = {f: f"V-{f}" for f in fields}
    values["个人简介"] = "首行\n次行"
    rg.merge_docx(path, out, values)
    from docx import Document
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert any("V-姓名" in t for t in texts)
    assert not any("{{" in t for t in texts)


def test_save_template_and_missing():
    path = _build_example()
    bytes_doc = rg.TEMPLATE_DIR.joinpath("sample.docx").read_bytes()
    info = rg.save_template("a.docx", bytes_doc)
    assert info["template_id"] and path
    assert rg.get_template("nope") is None
    tpl = rg.get_template(info["template_id"])
    assert "姓名" in (tpl.get("fields") or [])
    assert rg._missing(tpl["fields"], {"姓名": "张"}) == [f for f in tpl["fields"] if f != "姓名"]


def test_session_lifecycle_without_llm():
    tpl = rg.save_template("t.docx", _sample_bytes())
    s = rg.create_session(tpl["template_id"], None)
    assert s["status"] == "collecting"
    got = rg.get_session("missing")
    assert got is None
    # 直接 generate（无 AI 预填）也能产出文件
    out = rg.generate_docx(s["session_id"])
    assert out["file_name"] == "简历.docx"
    assert TEMPLATE_DIR.parent.joinpath("resume_gen", f"{s['session_id']}.docx").exists()


@pytest.mark.asyncio
async def test_chat_fills_and_done(monkeypatch):
    tpl = rg.save_template("t2.docx", _sample_bytes())
    s = rg.create_session(tpl["template_id"], None)
    fields_all = tpl["fields"]
    idx = {"n": 0}

    async def fake(prompt, system="", temperature=0.5):
        f = fields_all[idx["n"] % len(fields_all)]
        idx["n"] += 1
        return '{"fields":{"%s":"值%s"},"question":"还有吗"}' % (f, f)

    monkeypatch.setattr(rg, "llm_generate", fake)
    # 多轮直到 done
    turns = 0
    while True:
        r = await rg.chat(s["session_id"], "补充一句")
        turns += 1
        if r["done"] or turns > 30:
            break
    assert r["done"] is True
    assert r["missing"] == []
    got = rg.get_session(s["session_id"])
    assert got["status"] == "ready"


@pytest.mark.asyncio
async def test_chat_llm_malformed_raises(monkeypatch):
    tpl = rg.save_template("t3.docx", _sample_bytes())
    s = rg.create_session(tpl["template_id"], None)

    async def bad(*a, **k):
        return "not json at all"
    monkeypatch.setattr(rg, "llm_generate", bad)
    with pytest.raises(ValueError):
        await rg.chat(s["session_id"], "hi")


# ---------- router ----------

def test_resume_gen_gated(client):
    assert client.get("/api/v1/resume-gen/example-template").status_code == 403
    assert client.post("/api/v1/resume-gen/sessions", json={"template_id": "x"}).status_code == 403


def test_example_to_desktop_disabled_in_tests(client, enable_api):
    # 测试环境关闭桌面保存 => 500（带错误信息）
    r = client.post("/api/v1/resume-gen/example-template/to-desktop")
    assert r.status_code == 500


def test_upload_plain_resume_builtin_mode(client, enable_api):
    from docx import Document
    from io import BytesIO
    d = Document()
    d.add_paragraph("张三")
    d.add_paragraph("电话：13800000000")
    d.add_paragraph("工作经历")
    d.add_paragraph("在某某公司做过后端开发，负责接口与部署。")
    buf = BytesIO(); d.save(buf)
    r = client.post("/api/v1/resume-gen/upload-template",
                    files={"file": ("简历.docx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    assert r.status_code == 200
    body = r.json()
    assert body.get("mode") == "builtin"
    assert "工作经历" in (body.get("fields") or [])


def test_example_to_desktop_success(client, enable_api, tmp_path, monkeypatch):
    import app.services.resume_gen as rgm
    monkeypatch.setattr(rgm.settings, "resume_save_desktop", True)
    monkeypatch.setattr(rgm, "_desktop_dir", lambda: tmp_path)
    r = client.post("/api/v1/resume-gen/example-template/to-desktop")
    assert r.status_code == 200
    name = r.json()["desktop"]["filename"]
    assert name.endswith(".docx")
    assert (tmp_path / name).exists()


def test_full_router_flow(client, enable_api, monkeypatch):
    # 下载示例模板 -> 上传 -> 会话 -> chat(mock) -> generate -> download
    r = client.get("/api/v1/resume-gen/example-template")
    assert r.status_code == 200 and r.content[:2] == b"PK"

    r = client.post("/api/v1/resume-gen/upload-template",
                    files={"file": ("tmpl.docx", r.content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    assert r.status_code == 200
    tpl_id = r.json()["template_id"]

    async def fake(prompt, system="", temperature=0.5):
        return '{"fields":{"姓名":"张三"},"question":"继续？"}'
    monkeypatch.setattr(rg, "llm_generate", fake)

    r = client.post("/api/v1/resume-gen/sessions", json={"template_id": tpl_id, "resume_id": "b28a6c61da2f"})
    assert r.status_code == 200
    sid = r.json()["session_id"]

    r = client.post(f"/api/v1/resume-gen/sessions/{sid}/chat", json={"message": "你好"})
    assert r.status_code == 200
    assert r.json()["fields"].get("姓名") == "张三"

    r = client.get(f"/api/v1/resume-gen/sessions/{sid}")
    assert r.status_code == 200 and r.json()["session_id"] == sid

    r = client.post(f"/api/v1/resume-gen/sessions/{sid}/generate")
    assert r.status_code == 200 and r.json()["file_name"]

    r = client.get(f"/api/v1/resume-gen/sessions/{sid}/download")
    assert r.status_code == 200 and r.content[:2] == b"PK"


def test_resume_gen_errors(client, enable_api):
    assert client.get("/api/v1/resume-gen/sessions/zzz").status_code == 404
    # 模板不存在 -> 400
    r = client.post("/api/v1/resume-gen/sessions", json={"template_id": "no-such"})
    assert r.status_code == 400
    # 非 docx / 非 zip 头 -> 400
    assert client.post("/api/v1/resume-gen/upload-template",
                       files={"file": ("x.pdf", b"%PDF-1.4", "application/pdf")}).status_code == 400
    assert client.post("/api/v1/resume-gen/upload-template",
                       files={"file": ("x.docx", b"plain-text-no-zip", "application/octet-stream")}).status_code == 400


# ---------- 保版式：小节识别与按小节重写 ----------

def _table_resume() -> bytes:
    from docx import Document
    from io import BytesIO
    d = Document()
    t = d.add_table(rows=0, cols=2)
    def add(a, b):
        row = t.add_row()
        row.cells[0].text = a
        row.cells[1].text = b
    add("个人简历", "")
    add("姓名", "张凯")
    add("电话", "13800000000")
    add("教育背景", "")
    add("", "2017-2021 某大学 计算机")
    add("专业技能", "")
    add("", "Python、FastAPI")
    buf = BytesIO(); d.save(buf)
    return buf.getvalue()


def test_analyze_layout_table():
    tokens = rg.analyze_layout_bytes(_table_resume())
    assert "姓名" in tokens and "电话" in tokens
    assert "教育背景" in tokens and "专业技能" in tokens


def test_rewrite_layout_keeps_structure():
    from docx import Document
    from io import BytesIO
    data = _table_resume()
    src = TEMPLATE_DIR / "layout_t.docx"
    out = TEMPLATE_DIR / "layout_out.docx"
    src.write_bytes(data)
    vals = {"姓名": "李四", "电话": "13900000000", "教育背景": "2022 某大学\n主修CS", "专业技能": "Go、K8s"}
    rg.rewrite_layout_docx(src, out, vals)
    doc = Document(str(out))
    tbl = doc.tables[0]
    assert len(tbl.rows) == len(Document(BytesIO(data)).tables[0].rows)
    text = "\n".join(c.text for row in tbl.rows for c in row.cells)
    assert "李四" in text and "13900000000" in text and "Go、K8s" in text
    assert "教育背景" in text and "专业技能" in text  # 标题保留
    assert "张凯" not in text and "13800000000" not in text
    assert "Python、FastAPI" not in text


def test_rewrite_layout_keeps_unfilled_original():
    from docx import Document
    from io import BytesIO
    rg._ensure_dirs()
    data = _table_resume()
    src = TEMPLATE_DIR / "layout_t2.docx"; out = TEMPLATE_DIR / "layout_out2.docx"
    src.write_bytes(data)
    rg.rewrite_layout_docx(src, out, {"姓名": "王五"})  # 其它小节不给值 -> 应保留原文
    doc = Document(str(out))
    text = "\n".join(c.text for row in doc.tables[0].rows for c in row.cells)
    assert "王五" in text
    assert "2017-2021 某大学 计算机" in text  # 未给值小节原样保留


def test_analyze_and_rewrite_paragraph_resume():
    from docx import Document
    from io import BytesIO
    rg._ensure_dirs()
    d = Document()
    d.add_paragraph("工作经历")
    d.add_paragraph("在某厂做过后端开发，负责接口与部署")
    d.add_paragraph("电话：13800000000")
    buf = BytesIO(); d.save(buf)
    tokens = rg.analyze_layout_bytes(buf.getvalue())
    assert "工作经历" in tokens and "电话" in tokens
    src = TEMPLATE_DIR / "p_t.docx"; out = TEMPLATE_DIR / "p_out.docx"
    src.write_bytes(buf.getvalue())
    rg.rewrite_layout_docx(src, out, {"工作经历": "改写了工作内容", "电话": "1390000"})
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "工作经历" in text and "改写了工作内容" in text and "1390000" in text
    assert "某厂" not in text


def test_convert_doc_failure_returns_none(monkeypatch):
    import subprocess
    def boom(*a, **k):
        raise RuntimeError("no word")
    monkeypatch.setattr(subprocess, "run", boom)
    assert rg.convert_doc_to_docx_bytes(b"d0cfjunk") is None


def test_rewrite_textbox_style_inserts_after_heading():
    from docx import Document
    from io import BytesIO
    rg._ensure_dirs()
    d = Document()
    d.add_paragraph("工作经历")
    d.add_paragraph("教育背景")
    buf = BytesIO(); d.save(buf)
    src = TEMPLATE_DIR / "tb_src.docx"; out = TEMPLATE_DIR / "tb_out.docx"
    src.write_bytes(buf.getvalue())
    rg.rewrite_layout_docx(src, out, {"工作经历": "第一段内容\n第二行"})
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "工作经历" in texts and "教育背景" in texts
    assert any("第一段内容" in t for t in texts)


def test_upload_unstructured_docx_still_succeeds(client, enable_api):
    from docx import Document
    from io import BytesIO
    d = Document()
    d.add_paragraph("这是一段没有任何标签的普通文字内容而已。")
    buf = BytesIO(); d.save(buf)
    r = client.post("/api/v1/resume-gen/upload-template",
                    files={"file": ("无标签.docx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    assert r.status_code == 200
    assert r.json().get("mode") == "builtin"


# ---------- 内置多模板 + 照片位 + 预览 + 内置会话 ----------

def test_render_all_templates_have_photo():
    from app.services.resume_gen import BUILTIN_TEMPLATES, BUILTIN_FIELDS
    from docx import Document
    from docx.oxml.ns import qn
    from io import BytesIO
    vals = {f: f"示例{f}" for f in BUILTIN_FIELDS}
    for key, meta in BUILTIN_TEMPLATES.items():
        data = rg.render_builtin_docx(key, vals)
        assert data[:2] == b"PK"
        if meta.get("photo") == "1":
            doc = Document(BytesIO(data))
            all_text = "".join((t.text or "") for t in doc.element.body.iter(qn("w:t")))
            assert "照片" in all_text  # 照片位（可能在嵌套表/XML 内）


def test_preview_html_contains_accent():
    from app.services.resume_gen import BUILTIN_TEMPLATES
    for key, meta in BUILTIN_TEMPLATES.items():
        html = rg.preview_template_html(key)
        assert meta["accent"] in html and "照片" in html


def test_ensure_builtin_template_idempotent():
    tid1 = rg.ensure_builtin_template("classic")
    tid2 = rg.ensure_builtin_template("classic")
    assert tid1 == tid2
    tpl = rg.get_template(tid1)
    assert tpl["mode"] == "builtin" and tpl["builtin_key"] == "classic"
    assert "姓名" in tpl["fields"]


def test_router_session_with_template_key(client, enable_api):
    r = client.post("/api/v1/resume-gen/sessions", json={"template_key": "jade"})
    assert r.status_code == 200
    data = r.json()
    assert data.get("template_key") == "jade"
    assert data.get("session_id") and len(data.get("missing") or []) >= 10
    # 不存在的模板 key -> 400
    assert client.post("/api/v1/resume-gen/sessions", json={"template_key": "nope"}).status_code == 400


# ---------- HTML 版式 / 整份预览 ----------

def test_render_resume_html():
    html = rg.render_resume_html("editorial", {"姓名": "李四", "个人简介": "测试简介内容"}, photo_src="")
    assert "李四" in html and "测试简介内容" in html and "个人简介" in html
    assert "page" in html and "photo" in html or "•" in html
    with_photo = rg.render_resume_html("classic", {"姓名": "王五"}, photo_src="/x/photo")
    assert "photo" in with_photo and "x/photo" in with_photo


def test_router_session_preview(client, enable_api):
    r = client.post("/api/v1/resume-gen/sessions", json={"template_key": "editorial"})
    sid = r.json()["session_id"]
    prev = client.get(f"/api/v1/resume-gen/sessions/{sid}/preview")
    assert prev.status_code == 200 and "候选人" in prev.text and "editorial" in prev.text.lower() or "page" in prev.text
    # 未传照片 -> photo 404
    assert client.get(f"/api/v1/resume-gen/sessions/{sid}/photo").status_code == 404
    # 模板预览（示例内容 HTML）
    t = client.get("/api/v1/resume-gen/templates/sidebar/preview")
    assert t.status_code == 200 and "张凯" in t.text


def test_render_html_all_keys_no_error():
    from app.services.resume_gen import BUILTIN_TEMPLATES
    sample = {"姓名": "甲", "个人简介": "简介", "技能": "A\nB",
              "工作经历": "2020-2022 · 公司 岗位", "项目经历": "P · X", "教育经历": "U"}
    for key in BUILTIN_TEMPLATES:
        html = rg.render_resume_html(key, sample)
        assert "page" in html and "甲" in html

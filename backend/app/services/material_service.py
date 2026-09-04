from __future__ import annotations

import asyncio
import logging
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from pypdf import PdfReader

from ..config import settings
from ..database import generate_id, get_db, now_iso, row_to_dict
from ..ollama_client import ollama_client
from ..vector_store import vector_store

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}

# 资料默认分块参数（字符级）
_CHUNK_SIZE = 800
_CHUNK_OVERLAP = 120

_Doc = dict[str, str]  # {"path": "相对路径/文件名", "text": "正文"}


def _read_pdf(file_path: Path) -> str:
    parts: list[str] = []
    reader = PdfReader(str(file_path))
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            parts.append(page_text)
    return "\n".join(parts)


def _read_docx(file_path: Path) -> str:
    doc = DocxDocument(str(file_path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _read_plain(file_path: Path) -> str:
    for enc in ("utf-8", "gbk", "latin-1"):
        try:
            return file_path.read_text(encoding=enc)
        except (UnicodeDecodeError, OSError):
            continue
    return ""


def extract_text(file_content: bytes, filename: str) -> tuple[str, str]:
    """解析单份资料文本。返回 (text, ext)。不支持/解析失败抛 ValueError。"""
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise ValueError(f"不支持的文件格式: {suffix}，仅支持 "
                         + "/".join(sorted(ALLOWED_EXTENSIONS)))

    tmp = settings.upload_dir / f"{generate_id()}_material{suffix}"
    try:
        tmp.write_bytes(file_content)
        if suffix == ".pdf":
            text = _read_pdf(tmp)
        elif suffix == ".docx":
            text = _read_docx(tmp)
        else:
            text = _read_plain(tmp)
    finally:
        tmp.unlink(missing_ok=True)

    if not text.strip():
        raise ValueError(f"解析失败：无法提取文本内容（{filename}）")
    return text, suffix


def _chunk_text(text: str, size: int = _CHUNK_SIZE, overlap: int = _CHUNK_OVERLAP) -> list[str]:
    if len(text) <= size:
        return [text] if text.strip() else []
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + size)
        chunk = text[start:end]
        if chunk.strip() and len(chunk.strip()) > 20:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = end - overlap
    return chunks


def _get_row(material_id: str) -> dict[str, Any] | None:
    with get_db() as db:
        row = db.execute(
            "SELECT * FROM materials WHERE id = ?", (material_id,)
        ).fetchone()
        return row_to_dict(row)


def list_materials() -> list[dict[str, Any]]:
    with get_db() as db:
        rows = db.execute(
            """SELECT id, name, original_filename, file_ext, kind, size, index_status,
                      chunk_count, indexed_chunks, total_chunks, created_at
               FROM materials ORDER BY created_at DESC"""
        ).fetchall()
        return [row_to_dict(r) for r in rows]


async def upload_folder(folder_name: str, files: list[dict[str, Any]], *,
                        kind: str = "folder",
                        original_filename: str | None = None,
                        file_ext: str | None = None,
                        size: int = 0) -> dict[str, Any]:
    """把一个文件夹或单个文件导入为一个资料单元。

    files 内每项含 relpath(相对路径) 与 content(原始字节)；
    kind 区分 folder（整夹吸收）与 file（单个文件），单文件额外记录原始文件名/扩展名/大小。
    只吸收允许格式的文件；可解析文件不足一个时抛 ValueError。
    """
    if not folder_name.strip():
        raise ValueError("名称不能为空")

    docs: list[_Doc] = []
    skipped = 0
    for f in files:
        relpath = (f.get("relpath") or "").strip().replace("\\", "/")
        content = f.get("content") or b""
        if not relpath or not content:
            continue
        suffix = Path(relpath).suffix.lower()
        if suffix not in ALLOWED_EXTENSIONS:
            skipped += 1
            continue
        try:
            text, _ = extract_text(content, relpath)
        except ValueError:
            skipped += 1
            continue
        docs.append({"path": relpath, "text": text})

    if not docs:
        raise ValueError("没有可解析的 PDF/DOCX/TXT/MD 文件")

    name = folder_name.strip()
    original_filename = original_filename if original_filename is not None else name
    file_ext = file_ext if file_ext is not None else ("folder" if kind == "folder" else "")

    material_id = generate_id()
    with get_db() as db:
        db.execute(
            """INSERT INTO materials (id, name, original_filename, file_ext, kind, size, index_status)
               VALUES (?, ?, ?, ?, ?, ?, 'idle')""",
            (material_id, name, original_filename, file_ext, kind, size),
        )

    asyncio.create_task(_index_docs(material_id, docs))
    return {
        "material_id": material_id,
        "name": folder_name.strip(),
        "index_status": "processing",
        "file_count": len(docs),
        "skipped": skipped,
    }


def _scan_folder_files(root: Path) -> list[dict[str, Any]]:
    """递归读取一个文件夹里所有允许格式的文件（相对路径 + 字节），与代码库扫描一致。"""
    payload: list[dict[str, Any]] = []
    for fp in sorted(root.rglob("*")):
        if not fp.is_file():
            continue
        if fp.suffix.lower() not in ALLOWED_EXTENSIONS:
            continue
        rel = fp.relative_to(root).as_posix()
        if any(seg.startswith(".") for seg in Path(rel).parts):
            continue
        try:
            content = fp.read_bytes()
        except OSError:
            continue
        payload.append({"relpath": rel, "content": content})
    return payload


async def import_path(name: str, path_str: str) -> dict[str, Any]:
    """按本地路径导入为一个资料单元：文件夹则整夹吸收；单个文件则只导入该文件。"""
    path = Path(path_str)
    if path.is_dir():
        if not name.strip():
            name = path.name
        files = _scan_folder_files(path)
        if not files:
            raise ValueError("该文件夹内没有可解析的 PDF/DOCX/TXT/MD 文件")
        return await upload_folder(name, files)

    if path.is_file():
        suffix = path.suffix.lower()
        if suffix not in ALLOWED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {path.name}，仅支持 "
                             + "/".join(sorted(ALLOWED_EXTENSIONS)))
        try:
            content = path.read_bytes()
        except OSError:
            raise ValueError(f"无法读取文件: {path_str}")
        if not name.strip():
            name = path.stem
        return await upload_folder(
            name,
            [{"relpath": path.name, "content": content}],
            kind="file",
            original_filename=path.name,
            file_ext=suffix,
            size=len(content),
        )

    raise ValueError(f"路径不存在: {path_str}")


async def _index_docs(material_id: str, docs: list[_Doc]) -> None:
    meta = _get_row(material_id) or {}
    with get_db() as db:
        db.execute(
            "UPDATE materials SET index_status = 'processing' WHERE id = ?",
            (material_id,),
        )

    segments: list[tuple[str, str]] = []
    for doc in docs:
        text = (doc.get("text") or "").strip()
        path = (doc.get("path") or "").replace("\\", "/")
        if not text:
            continue
        for chunk in _chunk_text(text):
            segments.append((path, chunk))

    total = len(segments)
    with get_db() as db:
        db.execute(
            "UPDATE materials SET total_chunks = ? WHERE id = ?", (total, material_id)
        )

    try:
        if not segments:
            with get_db() as db:
                db.execute(
                    """UPDATE materials SET index_status = 'completed',
                           chunk_count = 0, total_chunks = 0 WHERE id = ?""",
                    (material_id,),
                )
            return

        ids = [f"material_{material_id}_{i}" for i in range(len(segments))]
        documents = [seg[1] for seg in segments]
        metadatas = [
            {
                "file_ext": meta.get("file_ext") or "",
                "title": meta.get("name", ""),
                "chunk_type": "doc",
                # 保留文件在资料内的相对路径，便于展示来源
                "file_path": seg[0] or meta.get("name", ""),
            }
            for seg in segments
        ]

        batch_size = 10
        for i in range(0, len(ids), batch_size):
            end = min(i + batch_size, len(ids))
            embeddings_data = await ollama_client.embed(documents[i:end])
            vector_store.add_chunks(
                material_id,
                ids[i:end],
                documents[i:end],
                embeddings_data,
                metadatas[i:end],
                collection_prefix="material_",
            )
            with get_db() as db:
                db.execute(
                    "UPDATE materials SET indexed_chunks = ? WHERE id = ?",
                    (end, material_id),
                )

        with get_db() as db:
            db.execute(
                """UPDATE materials SET index_status = 'completed',
                       chunk_count = ?, updated_at = ? WHERE id = ?""",
                (total, now_iso(), material_id),
            )
    except Exception as e:
        logger.exception("Material indexing failed for %s", material_id)
        with get_db() as db:
            db.execute(
                "UPDATE materials SET index_status = 'failed' WHERE id = ?",
                (material_id,),
            )
        raise


def rename_material(material_id: str, name: str) -> bool:
    with get_db() as db:
        db.execute(
            "UPDATE materials SET name = ?, updated_at = ? WHERE id = ?",
            (name, now_iso(), material_id),
        )
        return db.total_changes > 0


def delete_material(material_id: str) -> bool:
    vector_store.delete_collection(material_id, collection_prefix="material_")
    with get_db() as db:
        db.execute("DELETE FROM materials WHERE id = ?", (material_id,))
        return db.total_changes > 0

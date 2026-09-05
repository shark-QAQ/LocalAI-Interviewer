from __future__ import annotations

import asyncio
import hashlib
import json
import logging
import re
from pathlib import Path
from typing import Any

from pypdf import PdfReader
from docx import Document as DocxDocument

from ..config import settings
from ..database import generate_id, get_db, row_to_dict, now_iso
from .. import embed_client
from ..llm_client import llm_generate
from ..vector_store import vector_store
from .resume_code_map import ensure_confident

logger = logging.getLogger(__name__)

_TECH_SKILLS = [
    "Java", "Python", "JavaScript", "TypeScript", "Go", "Rust", "C", "C++", "C#",
    "Kotlin", "Scala", "Ruby", "PHP", "Swift", "Objective-C",
    "Spring", "Spring Boot", "Spring Cloud", "Django", "Flask", "FastAPI",
    "Express", "NestJS", "React", "Vue", "Angular", "Next.js",
    "Node.js", "Deno", "Bun",
    "MySQL", "PostgreSQL", "SQLite", "Oracle", "SQL Server", "MongoDB",
    "Redis", "Elasticsearch", "Memcached", "Neo4j", "InfluxDB",
    "Kafka", "RabbitMQ", "RocketMQ", "ActiveMQ",
    "Docker", "Kubernetes", "Helm", "Istio", "Nginx", "Apache",
    "AWS", "Azure", "GCP", "阿里云", "腾讯云",
    "Git", "Jenkins", "GitLab CI", "GitHub Actions", "Travis CI",
    "Linux", "Shell", "Bash",
    "TensorFlow", "PyTorch", "Scikit-learn", "Pandas", "NumPy",
    "OpenCV", "NLTK", "SpaCy", "Hugging Face", "LangChain",
    "GraphQL", "REST", "gRPC", "WebSocket", "Protobuf",
    "HTML", "CSS", "SASS", "LESS", "Tailwind",
    "Webpack", "Vite", "Rollup", "Babel",
    "RabbitMQ", "Celery", "gunicorn", "uvicorn",
    "Pytest", "JUnit", "TestNG", "Mocha", "Jest",
    "设计模式", "微服务", "分布式", "高并发", "高可用",
]


def _extract_text_from_pdf(file_path: Path) -> str:
    text_parts: list[str] = []
    reader = PdfReader(str(file_path))
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_text_from_docx(file_path: Path) -> str:
    doc = DocxDocument(str(file_path))
    text_parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)
    return "\n".join(text_parts)


def _extract_skills(text: str) -> list[str]:
    found: list[str] = []
    text_lower = text.lower()
    for skill in _TECH_SKILLS:
        if skill.lower() in text_lower:
            found.append(skill)
    return sorted(set(found))


def _extract_years(text: str) -> float | None:
    patterns = [
        r"(\d+\.?\d*)\s*年工作经验",
        r"(\d+\.?\d*)\s*年开发经验",
        r"(\d+\.?\d*)\s*years?\s*(?:of\s*)?experience",
        r"工作年限[：:\s]*(\d+\.?\d*)\s*年?",
        r"工作[：:\s]*(\d+\.?\d*)\s*年",
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            try:
                return float(m.group(1))
            except ValueError:
                continue
    return None


def _extract_name(text: str) -> str:
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    for line in lines[:10]:
        if re.match(r"^[\u4e00-\u9fa5]{2,5}$", line):
            return line
        if re.match(r"^[A-Z][a-z]+ [A-Z][a-z]+$", line):
            return line
    return "未知"


# 项目经历常见的“起止时间”标记，用于把一个大段里的多个项目切开
_PROJECT_DATE = re.compile(
    r"(?:19|20)\d{2}\s*[./\-年]?\s*\d{0,2}\s*(?:月)?\s*[-—–~至到]\s*"
    r"(?:(?:19|20)\d{2}\s*[./\-年]?\s*\d{0,2}|至今|今|present|now|current)",
    re.IGNORECASE,
)
# 看起来像“项目标题行”：短、无标点长句、无列表符号开头
_PROJECT_BULLET_START = re.compile(r"^\s*[•·▪●○◆\-—\d+\.、)]+\s*")


def _strip_project_date(line: str) -> str:
    return _PROJECT_DATE.sub("", line).strip(" |｜,，;；:：-\t")


def _looks_like_project_title(line: str) -> bool:
    s = line.strip()
    if not s or len(s) > 70 or _PROJECT_DATE.search(s):
        return False
    if _PROJECT_BULLET_START.match(s):
        return False
    # 排除明显是正文的句子（含逗号分号或较长描述）
    if any(c in s for c in "，。；：,.;:"):
        return False
    return True


def _split_projects(lines: list[str]) -> list[dict[str, str]]:
    """把“项目经历”段按起止时间行切成多个项目（兼容“标题行在前、时间行在后”）。"""
    if not lines:
        return []

    # 1) 用起止时间行定位每个项目的开头
    boundaries = [i for i, ln in enumerate(lines) if _PROJECT_DATE.search(ln)]

    def _make(header: str, body: list[str]) -> dict[str, str]:
        name = (header or "").strip(" |｜,，。;；:-")
        if not name:
            name = "项目"
        return {
            "name": name[:50],
            "description": "\n".join(body[:6])[:600] if body else "",
        }

    if boundaries:
        out: list[dict[str, str]] = []
        for k, bi in enumerate(boundaries):
            end = boundaries[k + 1] if k + 1 < len(boundaries) else len(lines)
            # 标题可能写在“时间行”紧邻的上一行
            header = _strip_project_date(lines[bi])
            if not header and bi > 0 and _looks_like_project_title(lines[bi - 1]):
                header = lines[bi - 1]
            body = [ln for ln in lines[bi + 1:end] if _PROJECT_DATE.search(ln) is None]
            out.append(_make(header, body))
        return out

    # 2) 兜底：没有时间行时，用“短标题行（如 1. / （公司））…”切分
    segs: list[list[str]] = []
    for ln in lines:
        s = ln.strip()
        starts_new = bool(re.match(r"^\d+[\.、)]\s*", s)) or (
            _looks_like_project_title(s) and "（" in s and "）" in s
        )
        if starts_new:
            segs.append([ln])
        elif segs:
            segs[-1].append(ln)
        else:
            segs.append([ln])
    if len(segs) >= 2:
        return [
            {"name": seg[0][:50], "description": "\n".join(seg[1:5])[:600]}
            for seg in segs
        ]
    # 3) 实在无法切分，退化成整段取首行作名字
    return [{"name": lines[0][:50], "description": "\n".join(lines[1:5])[:600]}]


def _extract_projects(text: str) -> list[dict[str, str]]:
    project_pattern = re.compile(
        r"(?:项目经验|项目经历|PROJECT)[：:\s]*\n(.*?)(?=\n(?:教育|工作|技能|自我|荣誉|证书|EDUCATION|WORK|SKILLS)|\Z)",
        re.DOTALL | re.IGNORECASE,
    )
    projects: list[dict[str, str]] = []
    for match in project_pattern.finditer(text):
        block = match.group(1).strip()
        lines = [l.strip() for l in block.split("\n") if l.strip()]
        projects.extend(_split_projects(lines))

    # 去重（同名的相邻重复项只留一次），最多返回 12 条
    seen: list[str] = []
    out: list[dict[str, str]] = []
    for p in projects:
        key = p["name"]
        if key not in seen:
            seen.append(key)
            out.append(p)
    return out[:12]


def parse_resume(file_path: Path) -> dict[str, Any]:
    with open(file_path, "rb") as f:
        header = f.read(8)

    if header[:3] == b'%PDF':
        text = _extract_text_from_pdf(file_path)
    elif header[:4] == b'PK\x03\x04':
        text = _extract_text_from_docx(file_path)
    else:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            text = _extract_text_from_pdf(file_path)
        elif suffix == ".docx":
            text = _extract_text_from_docx(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")

    if not text.strip():
        raise ValueError("解析失败：无法提取文本内容")

    skills = _extract_skills(text)
    years = _extract_years(text)
    name = _extract_name(text)
    projects = _extract_projects(text)

    parsed = {
        "name": name,
        "skills": skills,
        "years_of_experience": years,
        "projects": projects,
        "raw_length": len(text),
    }
    return {"text": text, "parsed": parsed, "skills": skills}


_RESUME_PARSE_SYSTEM = """你是一位专业的简历解析助手。请从以下简历文本中提取结构化信息。

请严格按 JSON 格式返回，包含以下字段：
{
  "name": "候选人姓名",
  "skills": ["技能1", "技能2", ...],
  "years_of_experience": 3.5,
  "projects": [{"name": "项目名", "description": "项目简介"}],
  "education": "学历信息（如有）",
  "summary": "一句话总结候选人背景"
}

注意：
- skills 只提取技术技能，不要包含软技能
- years_of_experience 为数字（年），无法判断时填 null
- projects 只提取简历中明确描述的项目
- 只返回 JSON，不要其他文字"""


async def parse_resume_with_llm(text: str) -> dict[str, Any]:
    """使用大模型解析简历文本，返回结构化数据。"""
    # 截取前 4000 字符避免超长
    truncated = text[:4000]
    raw = await llm_generate(
        prompt=f"请解析以下简历文本：\n\n{truncated}",
        system=_RESUME_PARSE_SYSTEM,
        temperature=0.1,
    )
    # 提取 JSON
    raw = (raw or "").strip()
    # 兼容 markdown code block
    if "```" in raw:
        import re as _re
        m = _re.search(r"```(?:json)?\s*\n?(.*?)\n?```", raw, _re.DOTALL)
        if m:
            raw = m.group(1).strip()
    try:
        parsed = json.loads(raw)
    except json.JSONDecodeError:
        logger.warning("LLM resume parse JSON decode failed, raw=%s", raw[:200])
        # fallback：用正则解析
        return parse_resume_fallback(text)

    skills = parsed.get("skills") or []
    name = parsed.get("name") or "未知"
    years = parsed.get("years_of_experience")
    projects = parsed.get("projects") or []

    result = {
        "name": name,
        "skills": skills,
        "years_of_experience": years,
        "projects": projects,
        "raw_length": len(text),
        "summary": parsed.get("summary", ""),
        "education": parsed.get("education", ""),
    }
    return {"text": text, "parsed": result, "skills": skills}


def parse_resume_fallback(text: str) -> dict[str, Any]:
    """正则解析兜底（LLM 解析失败时使用）。"""
    skills = _extract_skills(text)
    years = _extract_years(text)
    name = _extract_name(text)
    projects = _extract_projects(text)
    parsed = {
        "name": name,
        "skills": skills,
        "years_of_experience": years,
        "projects": projects,
        "raw_length": len(text),
    }
    return {"text": text, "parsed": parsed, "skills": skills}


def _chunk_resume(text: str, resume_id: str, parsed: dict) -> list[dict[str, Any]]:
    """将简历按模块分块"""
    chunks: list[dict[str, Any]] = []

    sections = re.split(r"\n(?=[\u4e00-\u9fa5]{2,8}[：:：])", text)

    for i, section in enumerate(sections):
        section = section.strip()
        if len(section) < 20:
            continue
        # detect section type
        section_type = "experience"
        if re.search(r"技能|技术栈|SKILLS", section):
            section_type = "skills"
        elif re.search(r"项目|PROJECT", section):
            section_type = "project"
        elif re.search(r"教育|EDUCATION", section):
            section_type = "education"

        # split long sections
        chunk_size = 800
        for j in range(0, len(section), chunk_size):
            chunk_text = section[j:j + chunk_size]
            if len(chunk_text.strip()) > 30:
                chunks.append({
                    "text": chunk_text,
                    "metadata": {
                        "resume_id": resume_id,
                        "chunk_type": section_type,
                        "section_index": i,
                    },
                })

    # add skills summary as a chunk
    if parsed.get("skills"):
        skill_chunk = f"候选人技能栈: {', '.join(parsed['skills'])}"
        if parsed.get("years_of_experience"):
            skill_chunk += f"\n工作年限: {parsed['years_of_experience']}年"
        chunks.append({
            "text": skill_chunk,
            "metadata": {
                "resume_id": resume_id,
                "chunk_type": "summary",
                "section_index": 999,
            },
        })

    return chunks


async def _index_resume(resume_id: str, text: str, parsed: dict) -> None:
    """将简历向量化并存入 ChromaDB"""
    try:
        with get_db() as db:
            db.execute("UPDATE resumes SET index_status = 'processing' WHERE id = ?", (resume_id,))

        chunks = _chunk_resume(text, resume_id, parsed)
        total = len(chunks)

        with get_db() as db:
            db.execute("UPDATE resumes SET total_chunks = ? WHERE id = ?", (total, resume_id))

        if not chunks:
            with get_db() as db:
                db.execute(
                    "UPDATE resumes SET index_status = 'completed', chunk_count = 0, total_chunks = 0 WHERE id = ?",
                    (resume_id,),
                )
            return

        ids = [f"resume_{resume_id}_{i}" for i in range(len(chunks))]
        documents = [c["text"] for c in chunks]
        metadatas = [c["metadata"] for c in chunks]

        batch_size = 10
        for i in range(0, len(ids), batch_size):
            end = min(i + batch_size, len(ids))
            batch_docs = documents[i:end]
            embeddings_data = await embed_client.embed(batch_docs)
            vector_store.add_chunks(
                resume_id,
                ids[i:end],
                batch_docs,
                embeddings_data,
                metadatas[i:end],
                collection_prefix="resume_",
            )
            with get_db() as db:
                db.execute("UPDATE resumes SET indexed_chunks = ? WHERE id = ?", (end, resume_id))

        with get_db() as db:
            db.execute(
                "UPDATE resumes SET index_status = 'completed', chunk_count = ?, updated_at = ? WHERE id = ?",
                (total, now_iso(), resume_id),
            )

    except Exception as e:
        logger.exception("Resume indexing failed for %s", resume_id)
        with get_db() as db:
            db.execute("UPDATE resumes SET index_status = 'failed' WHERE id = ?", (resume_id,))


def _persist_original(file_content: bytes, filename: str, file_hash: str) -> Path:
    """把上传的原始简历留在项目内 data/resumes/，便于后续改解析后重跑。"""
    d = settings.data_dir / "resumes"
    d.mkdir(parents=True, exist_ok=True)
    ext = Path(filename).suffix.lower() or ".pdf"
    p = d / f"{file_hash}{ext}"
    p.write_bytes(file_content)
    return p


async def upload_resume(
    file_content: bytes,
    filename: str,
    skip_vectorize: bool = False,
    use_llm: bool = False,
) -> dict[str, Any]:
    file_hash = hashlib.md5(file_content).hexdigest()
    with get_db() as db:
        existing = db.execute(
            "SELECT id FROM resumes WHERE raw_text_hash = ?", (file_hash,)
        ).fetchone()

    original = _persist_original(file_content, filename, file_hash)
    try:
        # 提取文本
        head = open(original, "rb").read(1024)
        is_pdf = b'%PDF' in head
        is_docx = head[:4] == b'PK\x03\x04'
        if is_pdf:
            text = _extract_text_from_pdf(original)
        elif is_docx:
            text = _extract_text_from_docx(original)
        else:
            suffix = original.suffix.lower()
            if suffix == ".pdf":
                text = _extract_text_from_pdf(original)
            elif suffix == ".docx":
                text = _extract_text_from_docx(original)
            else:
                raise ValueError(f"不支持的文件格式: {suffix}")

        if not text.strip():
            raise ValueError("解析失败：无法提取文本内容")

        # 根据模式选择解析方式
        if use_llm:
            result = await parse_resume_with_llm(text)
        else:
            result = parse_resume_fallback(text)
    except Exception:
        original.unlink(missing_ok=True)
        raise

    # 自动补全“简历项目 ↔ 代码库”映射（不打扰用户；匹配不到则留空按简历经历作答）
    try:
        with get_db() as db:
            code_repos = [
                {"name": r["name"]}
                for r in db.execute(
                    "SELECT name FROM projects WHERE index_status = 'completed' AND chunk_count > 0"
                ).fetchall()
            ]
        ensure_confident(result["parsed"].get("projects") or [], code_repos)
    except Exception:
        pass

    with get_db() as db:
        if existing:
            # 同文件重复上传：用新的（改进版）解析覆盖旧结果，而不是直接返回旧的
            resume_id = existing["id"]
            db.execute(
                """UPDATE resumes
                   SET candidate_name = ?, parsed_data = ?, skills = ?, index_status = 'idle',
                       updated_at = ?
                   WHERE id = ?""",
                (
                    result["parsed"]["name"],
                    json.dumps(result["parsed"], ensure_ascii=False),
                    json.dumps(result["skills"], ensure_ascii=False),
                    now_iso(),
                    resume_id,
                ),
            )
        else:
            resume_id = generate_id()
            db.execute(
                """INSERT INTO resumes (id, candidate_name, raw_text_hash, parsed_data, skills, index_status)
                   VALUES (?, ?, ?, ?, ?, 'idle')""",
                (
                    resume_id,
                    result["parsed"]["name"],
                    file_hash,
                    json.dumps(result["parsed"], ensure_ascii=False),
                    json.dumps(result["skills"], ensure_ascii=False),
                ),
            )

    # 向量化（可跳过）
    if skip_vectorize:
        with get_db() as db:
            db.execute(
                "UPDATE resumes SET index_status = 'completed', chunk_count = 0 WHERE id = ?",
                (resume_id,),
            )
        return {"resume_id": resume_id, "parsed_data": result["parsed"], "index_status": "completed"}

    # 解析逻辑升级后重传：作废旧向量，按新解析结果重建
    try:
        vector_store.delete_collection(resume_id, collection_prefix="resume_")
    except Exception:
        pass
    asyncio.create_task(_index_resume(resume_id, result["text"], result["parsed"]))

    return {"resume_id": resume_id, "parsed_data": result["parsed"], "index_status": "processing"}


def get_resume(resume_id: str) -> dict[str, Any] | None:
    with get_db() as db:
        row = db.execute("SELECT * FROM resumes WHERE id = ?", (resume_id,)).fetchone()
        if row is None:
            return None
        d = row_to_dict(row)
        d["parsed_data"] = json.loads(d["parsed_data"])
        d["skills"] = json.loads(d["skills"])
        return d


def list_resumes() -> list[dict[str, Any]]:
    with get_db() as db:
        rows = db.execute(
            "SELECT id, candidate_name, skills, index_status, chunk_count, created_at FROM resumes ORDER BY created_at DESC"
        ).fetchall()
        result = []
        for row in rows:
            d = row_to_dict(row)
            d["skills"] = json.loads(d["skills"]) if d.get("skills") else []
            result.append(d)
        return result


def delete_resume(resume_id: str) -> bool:
    vector_store.delete_collection(resume_id, collection_prefix="resume_")
    with get_db() as db:
        db.execute("DELETE FROM resumes WHERE id = ?", (resume_id,))
        return db.total_changes > 0


def rename_resume(resume_id: str, name: str) -> bool:
    with get_db() as db:
        db.execute(
            "UPDATE resumes SET candidate_name = ?, updated_at = ? WHERE id = ?",
            (name, now_iso(), resume_id),
        )
        return db.total_changes > 0

"""简历生成模块（docx 占位符模板 → AI 对话收集 → 生成 docx）。

持久化只用“项目内文件”，不引入数据库：
- 模板：data/resume_templates/<id>.docx + data/resume_templates/index.json
- 会话与字段状态：data/resume_gen/index.json（原子写）
- 产物：data/resume_gen/<session_id>.docx
满足“删项目即干净”，也无需 Docker/数据库服务。
"""

from __future__ import annotations

import json
import logging
import os
import re
import shutil
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.text.paragraph import Paragraph

from ..config import settings
from ..database import generate_id, now_iso
from ..llm_client import llm_generate

logger = logging.getLogger(__name__)

TEMPLATE_DIR: Path = settings.data_dir / "resume_templates"
GEN_DIR: Path = settings.data_dir / "resume_gen"
TEMPLATE_INDEX: Path = TEMPLATE_DIR / "index.json"
SESSION_INDEX: Path = GEN_DIR / "index.json"

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")

# 常见"标签：值"行的标签→占位符 映射（转模板时用来把真实值换成占位）
_INFO_LABELS: dict[str, str] = {
    "姓名": "姓名", "名字": "姓名", "name": "姓名",
    "求职意向": "求职意向", "意向岗位": "求职意向", "目标岗位": "求职意向",
    "电话": "电话", "手机": "电话", "手机号": "电话", "联系电话": "电话",
    "邮箱": "邮箱", "邮件": "邮箱", "email": "邮箱", "e-mail": "邮箱",
    "城市": "城市", "所在地": "城市", "现居": "城市", "居住地": "城市",
    "工作年限": "工作年限", "工作经验": "工作年限",
}
# 小节标题关键词 → 该节内容对应的占位符
_SECTION_KEYWORDS: list[tuple[str, str]] = [
    ("工作经历", "工作经历"), ("工作经验", "工作经历"),
    ("项目经历", "项目经历"), ("项目经验", "项目经历"),
    ("教育经历", "教育经历"), ("教育背景", "教育经历"),
    ("个人简介", "个人简介"), ("自我评价", "个人简介"), ("自我介绍", "个人简介"),
    ("专业技能", "技能"), ("技能专长", "技能"),
    ("求职意向", "求职意向"),
]


def _ensure_dirs() -> None:
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    GEN_DIR.mkdir(parents=True, exist_ok=True)


def _read_json(path: Path) -> dict[str, Any]:
    try:
        return json.loads(path.read_text(encoding="utf-8")) or {}
    except Exception:
        return {}


def _write_json(path: Path, data: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    try:
        tmp.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        os.replace(tmp, path)
    except Exception:
        try:
            tmp.unlink(missing_ok=True)
        except Exception:
            pass
        raise


# ---------------------------- docx 占位符工具 ----------------------------

def _iter_any_paragraph(doc: DocxDocument):
    """遍历整篇文档的每个 <w:p>（含表格单元格与文本框/形状内的段落），按文档顺序。"""
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph as DocParagraph
    for el in doc.element.body.iter(qn("w:p")):
        yield DocParagraph(el, doc)


def _iter_paragraphs(doc: DocxDocument):
    """递归遍历正文段落 + 表格内单元格段落 + 各节页眉/页脚段落。"""
    from docx.table import Table

    def walk_blocks(blocks: list) -> None:
        for b in blocks:
            if isinstance(b, Paragraph):
                yield b
            elif isinstance(b, Table):
                for row in b.rows:
                    for cell in row.cells:
                        yield from walk_blocks(cell.paragraphs)
                        for t in cell.tables:
                            yield from walk_blocks([t])

    yield from doc.paragraphs
    for t in doc.tables:
        yield from walk_blocks([t])
    for section in doc.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs
        for t in section.header.tables:
            yield from walk_blocks([t])
        for t in section.footer.tables:
            yield from walk_blocks([t])


def _scan_doc_fields(doc: DocxDocument) -> list[str]:
    """扫描 doc 里出现的 {{字段}}，按顺序去重。"""
    fields: list[str] = []
    seen: set[str] = set()
    for p in _iter_paragraphs(doc):
        for key in _PLACEHOLDER_RE.findall(p.text):
            key = key.strip()
            if key and key not in seen:
                seen.add(key)
                fields.append(key)
    return fields


def parse_template_fields(path: Path) -> list[str]:
    """扫描模板里的 {{字段}}，按出现顺序去重返回字段名。"""
    return _scan_doc_fields(DocxDocument(str(path)))


def scan_doc_fields_bytes(data: bytes) -> list[str]:
    """对 docx 字节内容扫占位字段（不落盘）。"""
    from io import BytesIO
    try:
        return _scan_doc_fields(DocxDocument(BytesIO(data)))
    except Exception:
        return []


# ---------------- 真实简历 → 占位简历（自动转换） ----------------

def _build_placeholder_doc() -> DocxDocument:
    """标准占位简历（字段齐全、无任何个人信息）。"""
    doc = DocxDocument()
    h = doc.add_paragraph()
    run = h.add_run("个人简历")
    run.bold = True
    run.font.size = __import__("docx.shared", fromlist=["Pt"]).Pt(20)
    h.alignment = 1

    def para(text: str, bold: bool = False) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold

    para("{{姓名}}", bold=True)
    para("求职意向：{{求职意向}}　电话：{{电话}}　邮箱：{{邮箱}}　城市：{{城市}}")
    para("个人简介", bold=True)
    para("{{个人简介}}")
    para("技能专长", bold=True)
    para("{{技能}}")
    para("工作经历", bold=True)
    para("{{工作经历}}")
    para("项目经历", bold=True)
    para("{{项目经历}}")
    para("教育经历", bold=True)
    para("{{教育经历}}")
    return doc


def _replace_in_paragraph(p: Paragraph, values: dict[str, str]) -> None:
    """把段落里的 {{key}} 换成 values[key]（未知占位保留原样），\n 用软换行保留；保留首 run 样式。"""
    text = p.text
    if "{{" not in text:
        return

    def repl(m: re.Match) -> str:
        key = m.group(1).strip()
        return str(values.get(key, m.group(0)))

    replaced = _PLACEHOLDER_RE.sub(repl, text)
    if replaced == text:
        return
    if not p.runs:
        return

    parts = replaced.split("\n")
    first = p.runs[0]
    first.text = parts[0]
    for part in parts[1:]:
        first.add_break()
        first.add_text(part)
    # 清空其余 run，避免旧占位内容残留
    for r in p.runs[1:]:
        r.text = ""


def merge_docx(template_path: Path, out_path: Path, values: dict[str, str]) -> None:
    """按 values 替换模板占位符并另存为 out_path（保留版式）。"""
    doc = DocxDocument(str(template_path))
    for p in _iter_paragraphs(doc):
        _replace_in_paragraph(p, values or {})
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def build_example_template(path: Path) -> None:
    """生成一份带 {{占位符}} 的标准占位简历模板（与自动转换用的结构一致）。"""
    path.parent.mkdir(parents=True, exist_ok=True)
    _build_placeholder_doc().save(str(path))


# ---------------- 内置专业模板（多套 + 照片位；出稿方向：布局可控，一定整齐） ----------------

BUILTIN_FIELDS = ["姓名", "求职意向", "电话", "邮箱", "城市", "个人简介", "技能", "工作经历", "项目经历", "教育经历"]

# key: {name, accent(hex), desc, style: line | bar | plain | sidebar | timeline}
BUILTIN_TEMPLATES: dict[str, dict[str, str]] = {
    "classic": {"name": "商务经典", "accent": "C23A2B", "style": "line", "photo": "1",
                "desc": "朱砂标题·深墨正文·顶部照片位"},
    "jade": {"name": "青玉现代", "accent": "5A7A6A", "style": "bar", "photo": "1",
             "desc": "青绿留白·标题带色条·顶部照片位"},
    "sidebar": {"name": "双栏侧边", "accent": "2F4858", "style": "sidebar", "photo": "1",
                "desc": "左深色栏·照片/联系/技能；右主经历"},
    "timeline": {"name": "时间轴", "accent": "2E7D6B", "style": "timeline", "photo": "0",
                 "desc": "经历左列时间·右列内容，清爽分列"},
    "editorial": {"name": "杂志大色块", "accent": "1F3A5F", "style": "editorial", "photo": "0",
                  "desc": "顶部深色大字块·杂志排版（差异最大）"},
    "modern": {"name": "现代商务", "accent": "3E5C9A", "style": "modern", "photo": "1",
               "desc": "蓝色标题带色条·清爽商务"},
    "minimal": {"name": "水墨极简", "accent": "211C17", "style": "plain", "photo": "1",
                "desc": "纯黑标题·极简素净·顶部照片位"},
}
BUILTIN_TEMPLATE_ORDER = ["classic", "editorial", "sidebar", "timeline", "modern", "minimal"]

_SECTION_TITLES = [
    ("个人简介", "个人简介"),
    ("技能专长", "技能"),
    ("工作经历", "工作经历"),
    ("项目经历", "项目经历"),
    ("教育经历", "教育经历"),
]


def _hex(color_hex: str) -> "RGBColor":
    from docx.shared import RGBColor
    h = (color_hex or "211C17").lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _cell_border(cell, color: str, sz: int = 6) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    tcB = tcPr.find(qn("w:tcBorders"))
    if tcB is None:
        tcB = OxmlElement("w:tcBorders"); tcPr.append(tcB)
    for edge in ("top", "left", "bottom", "right"):
        e = tcB.find(qn(f"w:{edge}"))
        if e is None:
            e = OxmlElement(f"w:{edge}"); tcB.append(e)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), color); e.set(qn("w:space"), "0")


def _photo_cell(cell, width_cm: float, height_cm: float, photo_path: str | None = None) -> None:
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl = cell.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(width_cm)
    pc = tbl.rows[0].cells[0]
    pc.width = Cm(width_cm)
    _cell_border(pc, "9A9A9A", 8)
    p = pc.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if photo_path:
        from pathlib import Path as _P
        if _P(photo_path).exists():
            run = p.add_run()
            run.add_picture(photo_path, width=Cm(width_cm - 0.4))
            return
    p.paragraph_format.space_before = Pt(height_cm * 22)
    p.paragraph_format.space_after = Pt(height_cm * 22)
    r = p.add_run("照片")
    r.font.name = "微软雅黑"
    r.font.size = Pt(10); r.font.color.rgb = _hex("BBBBBB")
    r.font.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def render_builtin_docx(template_key: str, values: dict[str, str], photo_path: str | None = None) -> bytes:
    """按内置模板（含照片位）生成 docx。布局可控，输出整齐。photo_path 可嵌入照片。"""
    from io import BytesIO
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    meta = BUILTIN_TEMPLATES.get(template_key) or BUILTIN_TEMPLATES["classic"]
    if meta["style"] == "sidebar":
        return _render_sidebar_docx(values, photo_path)
    if meta["style"] == "timeline":
        return _render_timeline_docx(values)
    accent = _hex(meta["accent"])
    style = meta["style"]
    ink = RGBColor(0x21, 0x1C, 0x17)
    gray = RGBColor(0x5A, 0x5F, 0x5C)

    v = {str(k): (str(val or "")) for k, val in (values or {}).items()}
    doc = DocxDocument()
    for s in doc.sections:
        s.top_margin = Cm(1.4); s.bottom_margin = Cm(1.4)
        s.left_margin = Cm(1.7); s.right_margin = Cm(1.7)

    def style_run(r, size: float, bold: bool, color):
        r.font.name = "微软雅黑"; r.font.size = Pt(size); r.bold = bold
        r.font.color.rgb = color
        r.font.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    def para(text="", size=10.5, bold=False, color=ink, align=None,
             space_after=3, space_before=0, indent_cm=0.0, line=1.22):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.line_spacing = line
        if indent_cm:
            p.paragraph_format.left_indent = Cm(indent_cm)
        if text:
            r = p.add_run(text); style_run(r, size, bold, color)
        return p

    # ---- 顶部：左侧姓名+联系行，右侧照片位 ----
    head = doc.add_table(rows=1, cols=2)
    head.autofit = False
    left, right = head.rows[0].cells
    left.width = Cm(13.6); right.width = Cm(3.0)
    name_p = left.paragraphs[0]
    nr = name_p.add_run(v.get("姓名") or "候选人")
    style_run(nr, 22, True, accent)
    contact = "   |   ".join(x for x in [
        ("意向：" + v["求职意向"]) if v.get("求职意向") else "",
        v.get("电话") or "", v.get("邮箱") or "", v.get("城市") or "",
    ] if x)
    if contact:
        cp = left.add_paragraph()
        cr = cp.add_run(contact); style_run(cr, 10.5, False, gray)
    _photo_cell(right, 2.7, 3.6, photo_path)

    def section_title(title: str) -> None:
        hp = para(title, 13, True, accent, space_after=3, space_before=7)
        pPr = hp._p.get_or_add_pPr()
        if style == "line":
            pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8")
            bottom.set(qn("w:color"), meta["accent"]); bottom.set(qn("w:space"), "2")
            pbdr.append(bottom); pPr.append(pbdr)
        elif style == "bar":
            pbdr = OxmlElement("w:pBdr"); leftb = OxmlElement("w:left")
            leftb.set(qn("w:val"), "single"); leftb.set(qn("w:sz"), "18")
            leftb.set(qn("w:color"), meta["accent"]); leftb.set(qn("w:space"), "4")
            pbdr.append(leftb); pPr.append(pbdr)
            hp.paragraph_format.left_indent = Cm(0.2)
        # plain：纯黑粗体即可

    def section(title: str, content: str) -> None:
        if not content:
            return
        section_title(title)
        for line in str(content).split("\n"):
            line = line.strip().lstrip("-·•* ").strip()
            if line:
                para("•  " + line, size=10.5, indent_cm=0.4)

    if v.get("个人简介"):
        section("个人简介", v["个人简介"])
    for title, key in _SECTION_TITLES:
        if key == "个人简介":
            continue
        if v.get(key):
            section(title, v[key])

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _run(doc_para, text, size=10.5, bold=False, color=None, name="微软雅黑") -> None:
    from docx.shared import Pt
    from docx.oxml.ns import qn
    r = doc_para.add_run(text)
    r.font.name = name; r.font.size = Pt(size); r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    r.font.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _shade(cell, fill_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _render_sidebar_docx(values: dict[str, str], photo_path: str | None = None) -> bytes:
    """双栏侧边：左深色栏（照片/联系/技能），右主经历。"""
    from io import BytesIO
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn

    v = {str(k): str(val or "") for k, val in (values or {}).items()}
    dark = _hex("2F4858"); white = RGBColor(0xFF, 0xFF, 0xFF)
    doc = DocxDocument()
    for s in doc.sections:
        s.top_margin = Cm(1.2); s.bottom_margin = Cm(1.2)
        s.left_margin = Cm(1.4); s.right_margin = Cm(1.4)

    def add_p(cell_or_doc, text="", size=10.5, bold=False, color=None, align=None, before=0, after=3):
        p = cell_or_doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(before)
        if text:
            _run(p, text, size=size, bold=bold, color=color)
        return p

    p = add_p(doc, v.get("姓名") or "候选人", 24, True, dark, WD_PARAGRAPH_ALIGNMENT.CENTER, after=2)
    contact = "   |   ".join(x for x in [
        ("意向：" + v["求职意向"]) if v.get("求职意向") else "",
        v.get("电话") or "", v.get("邮箱") or "", v.get("城市") or "",
    ] if x)
    if contact:
        add_p(doc, contact, 10, color=RGBColor(0x5A, 0x5F, 0x5C), align=WD_PARAGRAPH_ALIGNMENT.CENTER, after=8)

    body = doc.add_table(rows=1, cols=2)
    body.autofit = False
    left, right = body.rows[0].cells
    left.width = Cm(5.2); right.width = Cm(12.4)
    _shade(left, "2F4858")

    # 左栏：照片 + 联系 + 技能
    add_p(left, "", after=4)
    photo = left.add_table(rows=1, cols=1)
    photo.autofit = False
    pc = photo.rows[0].cells[0]; pc.width = Cm(3.6)
    _cell_border(pc, "FFFFFF", 8)
    if photo_path:
        from pathlib import Path as _P
        if _P(photo_path).exists():
            run = pc.paragraphs[0].add_run()
            run.add_picture(photo_path, width=Cm(3.2))
        else:
            photo_path = None
    if not photo_path:
        pp = add_p(pc, "照片", 10, color=RGBColor(0xE7, 0xE7, 0xE7), align=WD_PARAGRAPH_ALIGNMENT.CENTER)
        pp.paragraph_format.space_before = Pt(34); pp.paragraph_format.space_after = Pt(34)
    add_p(left, "联系方式", 13, True, white, before=14, after=4)
    for line in [v.get("电话") or "", v.get("邮箱") or "", v.get("城市") or ""]:
        if line:
            add_p(left, line, 10, color=white, after=2)
    skills = v.get("技能") or ""
    if skills:
        add_p(left, "技能专长", 13, True, white, before=12, after=4)
        for line in skills.split("\n"):
            line = line.strip().lstrip("-·•* ").strip()
            if line:
                add_p(left, "• " + line, 10, color=white, after=2)

    # 右栏：标题 + 简介 + 经历 + 教育
    def title(t):
        hp = add_p(right, t, 15, True, dark, before=2, after=3)
        pPr = hp._p.get_or_add_pPr()
        from docx.oxml import OxmlElement
        pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "10")
        bottom.set(qn("w:color"), "2F4858"); bottom.set(qn("w:space"), "2")
        pbdr.append(bottom); pPr.append(pbdr)

    def block(t, content):
        if not content:
            return
        title(t)
        for line in content.split("\n"):
            line = line.strip().lstrip("-·•* ").strip()
            if line:
                add_p(right, "•  " + line, 10.5, after=2)

    block("个人简介", v.get("个人简介") or "")
    block("工作经历", v.get("工作经历") or "")
    block("项目经历", v.get("项目经历") or "")
    block("教育经历", v.get("教育经历") or "")

    buf = BytesIO(); doc.save(buf)
    return buf.getvalue()


def _render_timeline_docx(values: dict[str, str]) -> bytes:
    """时间轴：经历每行 = 左列时间 + 右列内容（浅色网格行），观感与单栏明显不同。"""
    from io import BytesIO
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    v = {str(k): str(val or "") for k, val in (values or {}).items()}
    accent = _hex("2E7D6B")
    doc = DocxDocument()
    for s in doc.sections:
        s.top_margin = Cm(1.4); s.bottom_margin = Cm(1.4)
        s.left_margin = Cm(1.6); s.right_margin = Cm(1.6)

    def para(text, size=10.5, bold=False, color=None, align=None, after=3, indent=0.0):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(after)
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        if text:
            _run(p, text, size=size, bold=bold, color=color)
        return p

    para(v.get("姓名") or "候选人", 24, True, accent, WD_PARAGRAPH_ALIGNMENT.CENTER, after=2)
    contact = "   |   ".join(x for x in [
        ("意向：" + v["求职意向"]) if v.get("求职意向") else "",
        v.get("电话") or "", v.get("邮箱") or "", v.get("城市") or "",
    ] if x)
    if contact:
        para(contact, 10, color=RGBColor(0x5A, 0x5F, 0x5C), align=WD_PARAGRAPH_ALIGNMENT.CENTER, after=8)

    def title(t):
        p = para(t, 15, True, accent, after=3)
        p.paragraph_format.space_before = Pt(6)
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "10")
        bottom.set(qn("w:color"), "2E7D6B"); bottom.set(qn("w:space"), "2")
        pbdr.append(bottom); pPr.append(pbdr)

    def simple_block(t, content):
        if not content:
            return
        title(t)
        for line in content.split("\n"):
            line = line.strip().lstrip("-·•* ").strip()
            if line:
                para("•  " + line, indent=0.4)

    simple_block("个人简介", v.get("个人简介") or "")
    skills = v.get("技能") or ""
    if skills:
        title("技能专长")
        para("•  " + "　".join(l.strip() for l in skills.split("\n") if l.strip()), indent=0.4)

    # 经历用 时间|内容 网格行
    def timeline_block(t, content):
        if not content:
            return
        title(t)
        lines = [l.strip() for l in content.split("\n") if l.strip()]
        tbl = doc.add_table(rows=0, cols=2)
        tbl.autofit = False
        for ln in lines:
            time_part, _, rest = ln.partition("·")
            if not rest:
                rest, time_part = ln, ""
            row = tbl.add_row()
            row.cells[0].width = Cm(3.4); row.cells[1].width = Cm(12.6)
            if time_part:
                tp = row.cells[0].paragraphs[0]
                _run(tp, time_part.strip(), 10, True, accent)
            rp = row.cells[1].paragraphs[0]
            _run(rp, rest.strip(), 10.5)

    timeline_block("工作经历", v.get("工作经历") or "")
    timeline_block("项目经历", v.get("项目经历") or "")
    timeline_block("教育经历", v.get("教育经历") or "")

    buf = BytesIO(); doc.save(buf)
    return buf.getvalue()


def render_resume_html(template_key: str, values: dict[str, str], photo_src: str = "") -> str:
    """真正可打印的 HTML 版式（每套 CSS 结构差异大），浏览器可另存 PDF。

    结构：顶部姓名区 + 可选照片 <img>（src 用服务端照片地址）；分节列表按模板排版。
    """
    import html as _html

    v = {str(k): _html.escape(str(val or "")) for k, val in (values or {}).items()}
    key = template_key
    accent = BUILTIN_TEMPLATES.get(key, {}).get("accent", "C23A2B")
    name = v.get("姓名") or "候选人"
    contact = "   ·   ".join(x for x in [
        ("意向：" + v["求职意向"]) if v.get("求职意向") else "",
        v.get("电话") or "", v.get("邮箱") or "", v.get("城市") or "",
    ] if x)
    photo_block = f'<div class="photo"><img src="{_html.escape(photo_src, quote=True)}" alt=""/></div>' if photo_src else ""

    def bullets(t, content):
        if not content:
            return ""
        items = "".join(f"<li>{_html.escape(x.strip().lstrip('-·•* '))}</li>"
                        for x in content.split("\n") if x.strip())
        return f'<h3>{t}</h3><ul>{items}</ul>'

    sections = "".join([
        bullets("个人简介", v.get("个人简介") or ""),
        bullets("技能专长", v.get("技能") or ""),
        bullets("工作经历", v.get("工作经历") or ""),
        bullets("项目经历", v.get("项目经历") or ""),
        bullets("教育经历", v.get("教育经历") or ""),
    ])

    # 六套差异 CSS：按 key 决定背景/配色/栏式/标题样式
    css_common = (
        "*{box-sizing:border-box;margin:0;padding:0} body{font-family:'Microsoft YaHei',sans-serif;color:#211C17;"
        "background:#eef0f2;display:flex;justify-content:center;padding:26px}"
        ".page{width:760px;background:#fff;min-height:1000px;padding:46px 54px;box-shadow:0 4px 24px rgba(0,0,0,.12)}"
        "h3{margin:20px 0 8px;font-size:16px;letter-spacing:1px} ul{list-style:none;margin:0;padding:0}"
        "li{line-height:1.85;font-size:13.5px;margin:2px 0;padding-left:14px;position:relative}"
        "li:before{content:'•';position:absolute;left:0;color:#{accent}}"
    )
    styles = {
        "classic": (f"{css_common}"
            ".head{display:flex;align-items:center;gap:18px;border-bottom:3px solid #{accent};padding-bottom:12px}"
            ".name{font-size:28px;font-weight:800;color:#{accent}} .contact{color:#5A5F5C;font-size:12px;margin-top:4px}"
            ".photo{width:92px;height:120px;border:1px dashed #bbb;overflow:hidden;margin-left:auto}"
            ".photo img{width:100%;height:100%;object-fit:cover}"
            "h3{color:#{accent};border-bottom:1px solid #eee;padding-bottom:3px}"),
        "jade": (f"{css_common}"
            ".head{text-align:center;border-bottom:2px solid #{accent};padding-bottom:10px}"
            ".name{font-size:26px;font-weight:800;color:#{accent}} .contact{color:#5A5F5C;font-size:12px;margin-top:4px}"
            ".photo{width:92px;height:120px;margin:0 auto 10px;border:1px dashed #bbb;overflow:hidden}"
            ".photo img{width:100%;height:100%;object-fit:cover}"
            "h3{color:#{accent}}"),
        "sidebar": (f"{css_common}"
            ".head{padding-bottom:12px;border-bottom:3px solid #{accent};margin-bottom:14px}"
            ".name{font-size:26px;font-weight:800;color:#{accent}} .contact{color:#5A5F5C;font-size:12px}"
            ".cols{display:flex;gap:26px} .side{width:200px;flex-shrink:0} .main{flex:1}"
            ".side h3{color:#2F4858;border-bottom:2px solid #2F4858} h3{color:#2F4858}"),
        "timeline": (f"{css_common}"
            ".head{text-align:center;padding-bottom:10px;border-bottom:2px solid #{accent}}"
            ".name{font-size:26px;font-weight:800;color:#{accent}} .contact{color:#5A5F5C;font-size:12px}"
            "h3{color:#{accent}}"),
        "editorial": (f"{css_common}"
            ".head{background:#{accent};color:#fff;margin:-46px -54px 24px;padding:34px 54px}"
            ".head .name{font-size:34px;font-weight:900} .head .contact{color:rgba(255,255,255,.85);font-size:12px;margin-top:6px}"
            "h3{color:#{accent};font-size:18px}"),
        "modern": (f"{css_common}"
            ".head{display:flex;align-items:center;gap:16px;border-bottom:4px solid #{accent};padding-bottom:10px}"
            ".name{font-size:26px;font-weight:800;color:#{accent}} .contact{color:#5A5F5C;font-size:12px}"
            "h3{color:#{accent};font-size:15px;border-left:5px solid #{accent};padding-left:8px}"),
    }
    style = styles.get(key, styles["classic"])
    ac = accent
    e = _html.escape

    def lis(content):
        return "".join("<li>" + e(x.strip().lstrip("-·•* ")) + "</li>"
                       for x in content.split("\n") if x.strip())

    def sec(t, c):
        return (f"<h3>{e(t)}</h3><ul>{lis(c)}</ul>") if c else ""

    intro = sec("个人简介", v.get("个人简介"))
    skills = sec("技能专长", v.get("技能"))
    exp = sec("工作经历", v.get("工作经历")) + sec("项目经历", v.get("项目经历"))
    edu = sec("教育经历", v.get("教育经历"))

    if key == "sidebar":
        side_items = "".join("<li>" + e(x) + "</li>" for x in
                             [v.get("电话"), v.get("邮箱"), v.get("城市")] if x)
        inner = (
            f"<div style='display:flex;gap:26px;margin-top:14px'>"
            f"<div style='width:210px;flex-shrink:0;background:#2F4858;color:#fff;border-radius:6px;padding:18px 14px'>"
            f"{photo_block}"
            f"<h3 style='color:#fff'>联系</h3><ul>{side_items}</ul>{skills.replace('h3', 'h3 style=\"color:#fff\"')}"
            f"</div>"
            f"<div style='flex:1'>{intro}{exp}{edu}</div></div>"
        )
        shell_head = f"<div class='head'><div class='name'>{e(name)}</div><div class='contact'>{e(contact)}</div></div>"
    elif key == "editorial":
        chips = "".join(f"<span style='background:#eee;border-radius:12px;padding:2px 10px;margin:0 6px 6px 0;display:inline-block;font-size:12px'>{e(x)}</span>"
                        for x in [s.strip() for s in (v.get("技能") or "").replace("\n", "、").split("、") if s.strip()][:12])
        inner = (
            f"<div style='background:#{ac};color:#fff;margin:-46px -54px 22px;padding:34px 54px;border-radius:0 0 18px 18px'>"
            f"<div style='font-size:34px;font-weight:900;letter-spacing:2px'>{e(name)}</div>"
            f"<div style='opacity:.85;font-size:12px;margin-top:8px'>{e(contact)}</div></div>"
            f"<div style='margin-bottom:14px'>{chips}</div>{intro}{exp}{edu}"
        )
        shell_head = ""
    elif key == "timeline":
        rows = ""
        for field, label in (("工作经历", "工作"), ("项目经历", "项目"), ("教育经历", "教育")):
            content = v.get(field)
            if not content:
                continue
            rows += f"<h3>{e(label)}经历</h3>"
            for line in content.split("\n"):
                time_part, _, rest = line.partition("·")
                if not rest:
                    rest, time_part = line, ""
                rows += (
                    f"<div style='display:flex;gap:12px;margin:2px 0;'>"
                    f"<div style='width:150px;flex-shrink:0;font-weight:700;color:#{ac}'>{e(time_part.strip())}</div>"
                    f"<div style='flex:1'>{e(rest.strip())}</div></div>"
                )
        inner = (f"<div class='head' style='text-align:center'><div class='name'>{e(name)}</div>"
                 f"<div class='contact'>{e(contact)}</div></div>{intro}{skills}<div style='margin-top:10px'>{rows}</div>")
        shell_head = ""
    else:
        inner = f"{skills}{intro}{exp}{edu}"
        shell_head = (f"<div class='head'>{photo_block}<div><div class='name'>{e(name)}</div>"
                      f"<div class='contact'>{e(contact)}</div></div></div>")

    return (
        f"<!doctype html><html lang='zh'><head><meta charset='utf-8'>"
        f"<style>{style}</style></head><body><div class='page'>"
        f"{shell_head}{inner}</div></body></html>"
    )


def preview_template_html(template_key: str) -> str:
    """近似版式预览（与 docx 出稿视觉一致，供前端 iframe 预览）。"""
    meta = BUILTIN_TEMPLATES.get(template_key) or BUILTIN_TEMPLATES["classic"]
    accent = meta["accent"]
    title_dec = "border-bottom:2px solid #" + accent if meta["style"] == "line" else \
                "border-left:4px solid #" + accent + ";padding-left:8px" if meta["style"] == "bar" else ""
    demo = {
        "姓名": "张三", "求职意向": "AI 大模型应用工程师",
        "电话": "138-0000-0000", "邮箱": "zk@example.com", "城市": "上海",
    }
    contact = "   |   ".join([x for x in [
        "意向：" + demo["求职意向"], demo["电话"], demo["邮箱"], demo["城市"]]])
    sections = [
        ("个人简介", "五年后端与 AI 应用经验，聚焦 RAG 与大模型工程化。"),
        ("技能专长", "Python / FastAPI / RAG / LangGraph / Docker"),
        ("工作经历", "高级后端工程师 · 负责 AI 研判平台架构与工具链接入"),
        ("项目经历", "安全运营平台 · 特征工程 + 规则兜底 + LangGraph 工作流"),
        ("教育经历", "某大学 计算机科学与技术 · 本科"),
    ]
    body = "".join(
        f'<div style="color:#{accent};font-weight:700;font-size:15px;margin:16px 0 4px;{title_dec}">{t}</div>'
        f'<div style="font-size:13px;color:#333;line-height:1.7">{"<br/>".join("•  " + l for l in c.split(chr(10)))}</div>'
        for t, c in sections
    )
    return f"""<!doctype html><html lang="zh"><head><meta charset="utf-8">
<style>body{{font-family:'Microsoft YaHei';margin:0;padding:28px;background:#f7f5f0;color:#211C17}}
.paper{{max-width:760px;margin:0 auto;background:#fff;padding:34px 38px;box-shadow:0 2px 12px rgba(0,0,0,.06);border-radius:2px}}
.head{{display:flex;justify-content:space-between;align-items:center}}
.name{{font-size:26px;font-weight:800;color:#{accent}}}
.contact{{color:#5A5F5C;font-size:12px;margin-top:6px}}
.photo{{width:88px;height:118px;border:1px dashed #bbb;color:#bbb;display:flex;align-items:center;justify-content:center;font-size:12px;margin-left:16px}}</style></head>
<body><div class="paper">
<div class="head"><div><div class="name">{demo["姓名"]}</div>
<div class="contact">{contact}</div></div><div class="photo">照片</div></div>
{body}
</div></body></html>"""


def ensure_builtin_template(template_key: str) -> str:
    """确保某内置模板已注册，返回其 template_id（生成走内置渲染，不读上传文件）。"""
    if template_key not in BUILTIN_TEMPLATES:
        raise ValueError("未知模板：" + template_key)
    idx = _read_json(TEMPLATE_INDEX)
    for tid, meta in idx.items():
        if meta.get("mode") == "builtin" and meta.get("builtin_key") == template_key:
            return tid
    tid = generate_id()
    idx[tid] = {
        "file_name": f"内置模板-{BUILTIN_TEMPLATES[template_key]['name']}.docx",
        "fields": list(BUILTIN_FIELDS),
        "mode": "builtin",
        "builtin_key": template_key,
        "created_at": now_iso(),
    }
    _write_json(TEMPLATE_INDEX, idx)
    return tid


# ---------------- 保版式：识别“小节/信息标签”，按小节重写正文 ----------------

_EXTRA_INFO = {
    "出生年月": "出生年月", "民族": "民族", "身高": "身高", "性别": "性别",
    "学历": "学历", "毕业院校": "毕业院校", "年龄": "年龄",
}
_INFO_LOOKUP = {re.sub(r"\s+", "", k): v for k, v in {**_INFO_LABELS, **_EXTRA_INFO}.items()}
_SECTION_CANON = [
    "教育背景", "专业技能", "工作经历", "项目经历", "自我评价", "个人简介",
    "实习经历", "校园经历", "证书荣誉", "兴趣爱好", "个人荣誉",
]


def _clean(t: str) -> str:
    return re.sub(r"\s+", "", t or "")


def _kind_label(text: str) -> tuple[str, str] | None:
    """判断一段文本是(info标签/section标题)或“标签：值”行。返回 (kind, token) 或 None。"""
    c = _clean(text).rstrip("：:").rstrip()
    # “标签：值”行（如 电话：138…）
    m = re.split(r"[：:]", c, maxsplit=1)
    if len(m) == 2 and m[0] in _INFO_LOOKUP:
        return ("info", _INFO_LOOKUP[m[0]])
    for sec in _SECTION_CANON:
        if c == sec or (c.startswith(sec) and len(c) <= len(sec) + 1):
            return ("section", sec)
    if c in _INFO_LOOKUP:
        return ("info", _INFO_LOOKUP[c])
    return None


def _row_unique_cells(row):
    out, seen = [], set()
    for cell in row.cells:
        if id(cell) not in seen:
            seen.add(id(cell))
            out.append(cell)
    return out


def _cell_lines(cell, text: str) -> None:
    """把 cell 内容替换为 text（多行用软换行），保留该 cell 首段样式。"""
    if not cell.paragraphs:
        return
    first = cell.paragraphs[0]
    if first.runs:
        parts = text.split("\n")
        first.runs[0].text = parts[0]
        for part in parts[1:]:
            first.runs[0].add_break()
            first.runs[0].add_text(part)
        for r in first.runs[1:]:
            r.text = ""
    else:
        first.add_run(text)
    for p in cell.paragraphs[1:]:
        if p.runs:
            for r in p.runs:
                r.text = ""


def convert_doc_to_docx_bytes(data: bytes) -> bytes | None:
    """用本机 Word 把老式 .doc 转成 .docx（尽力而为；没装 Word 返回 None）。"""
    import subprocess
    import tempfile
    try:
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "in.doc"
            dst = Path(td) / "out.docx"
            src.write_bytes(data)
            script = (
                "$ErrorActionPreference='Stop'\n"
                "$w = New-Object -ComObject Word.Application\n"
                "try { $w.Visible=$false; $d=$w.Documents.Open('" + str(src) + "');"
                " $d.SaveAs2('" + str(dst) + "', 12); $d.Close() }"
                " finally { $w.Quit() }"
            )
            r = subprocess.run(
                ["powershell", "-NoProfile", "-NonInteractive", "-Command", script],
                capture_output=True, timeout=120,
            )
            if r.returncode == 0 and dst.exists():
                return dst.read_bytes()
    except Exception:
        return None
    return None


def _doc_has_table_text(doc: DocxDocument) -> bool:
    try:
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    if _clean(cell.text):
                        return True
    except Exception:
        return False
    return False


def analyze_layout_bytes(data: bytes) -> list[str]:
    """按出现顺序找出这份版式里需要“重写”的字段（标签 + 分节标题）。

    表格版式：逐行扫描“标签→值/空槽”与分节标题；文本框/普通段落版式：
    遍历整篇所有 <w:p>（含文本框内）识别标题与标签。识别出的都算待填字段。
    """
    from io import BytesIO
    doc = DocxDocument(BytesIO(data))
    tokens: list[str] = []
    seen: set[str] = set()

    def add(t: str) -> None:
        if t and t not in seen:
            seen.add(t)
            tokens.append(t)

    if _doc_has_table_text(doc):
        for table in doc.tables:
            for row in table.rows:
                cells = _row_unique_cells(row)
                kinds = [_kind_label(c.text) for c in cells]
                # 分节标题行
                for k in kinds:
                    if k and k[0] == "section":
                        add(k[1])
                        break
                # “标签 → 空/值槽”行（可能有多个标签连续）
                i = 0
                while i < len(cells):
                    k = kinds[i]
                    if k and k[0] == "info":
                        add(k[1])
                        i += 1
                    else:
                        i += 1
    else:
        for p in _iter_any_paragraph(doc):
            t = (p.text or "").strip()
            if not t:
                continue
            k = _kind_label(t)
            if not k:
                continue
            add(k[1])
    return tokens


def _set_para_text(p, text: str) -> None:
    """把段落文本替换为 text（多行软换行），保留首 run 样式。"""
    if not getattr(p, "runs", None):
        p.add_run(text)
        return
    parts = text.split("\n")
    p.runs[0].text = parts[0]
    for part in parts[1:]:
        p.runs[0].add_break()
        p.runs[0].add_text(part)
    for r in p.runs[1:]:
        r.text = ""


def _insert_after_paragraph(p, text: str) -> None:
    """在段落 p 之后插入一段新文本（用于文本框/无槽位版式的“插入式填充”）。"""
    from docx.oxml import OxmlElement
    el = OxmlElement("w:p")
    p._p.addnext(el)
    from docx.text.paragraph import Paragraph as DocPara
    newp = DocPara(el, p._parent if getattr(p, "_parent", None) else None)
    newp.add_run(text)
    return newp


def rewrite_layout_docx(template_path: Path, out_path: Path, values: dict[str, str]) -> None:
    """在原版式上重写：标题/标签/结构保留。

    - 表格版式：标签→值槽精确回填（空值槽也填），分节标题后的正文行替换/清空；
    - 文本框/普通段落版式：内容行原位替换；无槽位的标题/标签则“在标题后插入”新段。
    values 未给的字段一律保持原样。
    """
    doc = DocxDocument(str(template_path))
    vals = {str(k): str(v) for k, v in (values or {}).items() if v is not None}
    out_path.parent.mkdir(parents=True, exist_ok=True)

    if _doc_has_table_text(doc):
        for table in doc.tables:
            cur: str | None = None
            placed_sec: set[str] = set()
            for row in table.rows:
                cells = _row_unique_cells(row)
                kinds = [_kind_label(c.text) for c in cells]
                texts = [_clean(c.text) for c in cells]

                # 分节标题行
                new_sec = None
                for k in kinds:
                    if k and k[0] == "section":
                        new_sec = k[1]
                        break
                if new_sec:
                    cur = new_sec
                    continue

                # 信息对：标签 → 紧邻“非标签槽位”（允许空）回填
                i = 0
                while i < len(cells):
                    k = kinds[i]
                    if k and k[0] == "info" and k[1] in vals:
                        j = i + 1
                        while j < len(cells) and kinds[j] and kinds[j][0] in ("info", "section"):
                            j += 1
                        if j < len(cells):
                            _cell_lines(cells[j], vals[k[1]])
                        i = j
                        continue
                    i += 1

                # 分节正文：第一处非标签槽位填该节内容，其余清空
                content_idx = [
                    idx for idx, (kk, tt) in enumerate(zip(kinds, texts))
                    if tt and not (kk and kk[0] in ("info", "section"))
                ]
                if cur and cur in vals and content_idx:
                    if cur not in placed_sec:
                        placed_sec.add(cur)
                        main = max(content_idx, key=lambda idx: len(texts[idx]))
                        _cell_lines(cells[main], vals[cur])
                    else:
                        for idx in content_idx:
                            _cell_lines(cells[idx], "")
    else:
        # 文本框/普通段落版式：先原位替换，无槽位的标签/标题再插入
        paras = list(_iter_any_paragraph(doc))
        cur: str | None = None
        placed: set[str] = set()
        for p in paras:
            t = (p.text or "").strip()
            if not t:
                continue
            k = _kind_label(t)
            if k and k[0] == "section":
                cur = k[1]
                continue
            if k and k[0] == "info":
                parts = re.split(r"[：:]\s*", t, maxsplit=1)
                if len(parts) == 2:
                    label = _clean(parts[0])
                    tok = _INFO_LOOKUP.get(label)
                    if tok and tok in vals:
                        _set_para_text(p, f"{parts[0].strip()}：{vals[tok]}")
                        placed.add(tok)
                elif k[1] in vals:
                    # 仅标签无值：原位改为 标签：值
                    _set_para_text(p, f"{t}：{vals[k[1]]}")
                    placed.add(k[1])
                continue
            if cur:
                if cur in vals and cur not in placed:
                    placed.add(cur)
                    _set_para_text(p, vals[cur])
                elif cur in vals:
                    _set_para_text(p, "")  # 后续内容清空防重复
        # 插入式填充：正文/标签没有被替换到的，在所在段后补一段
        for p in paras:
            t = (p.text or "").strip()
            if not t:
                continue
            k = _kind_label(t)
            tok = k[1] if k else None
            if tok and tok in vals and tok not in placed and k[0] == "section":
                _insert_after_paragraph(p, vals[tok])
                placed.add(tok)
    doc.save(str(out_path))


# ---------------------------- 模板注册 ----------------------------

def register_template(
    file_name: str, file_bytes: bytes, fields: list[str], mode: str
) -> dict[str, Any]:
    """注册模板：mode='placeholder'（占位替换）| 'sections'（按小节重写，保版式）。"""
    _ensure_dirs()
    template_id = generate_id()
    target = TEMPLATE_DIR / f"{template_id}.docx"
    target.write_bytes(file_bytes)
    idx = _read_json(TEMPLATE_INDEX)
    idx[template_id] = {
        "file_name": file_name, "fields": fields, "mode": mode, "created_at": now_iso(),
    }
    _write_json(TEMPLATE_INDEX, idx)
    return {"template_id": template_id, "fields": fields, "file_name": file_name, "mode": mode}


def save_template(file_name: str, file_bytes: bytes) -> dict[str, Any]:
    """占位模板（含 {{}}）的注册入口（兼容测试/旧逻辑）。"""
    target_mode_file = TEMPLATE_DIR / f"_chk_{generate_id()}.docx"
    target_mode_file.write_bytes(file_bytes)
    try:
        fields = parse_template_fields(target_mode_file)
    finally:
        target_mode_file.unlink(missing_ok=True)
    return register_template(file_name, file_bytes, fields, "placeholder")


def get_template(template_id: str) -> dict[str, Any] | None:
    return _read_json(TEMPLATE_INDEX).get(template_id)


# ---------------------------- 会话（AI 收集） ----------------------------

_SYSTEM = """你是资深 HR/简历写作顾问。配合“简历模板的各填写小节”，把用户用自然语言讲的内容加工成可直接放进简历的正文。

规则：
1. 只输出一个 JSON 对象：{"fields": {小节名: 值}, "question": "你对用户说的下一句话"}。
2. 【务必润色】fields 的值必须是书面简历语，不是聊天原文：
   · 去掉口语、语气词、重复与客套（如“我这个人”“其实”“然后就是”）；
   · 工作/项目经历写成结构化：职责+做法+技术点+结果/量化（有数据就带数据，如“QPS 提升 40%”）；
   · 技能用顿号归类列出；个人简介 2~3 句概括定位与亮点。
3. 参考“已有简历素材”（你上传的简历）：沿用其经历细节、量化数字与专业措辞来润色；用户这次新说的信息优先采纳、与素材一致。
4. 素材没有且用户也没说的硬信息（公司名/数字/证书）不得编造；实在缺的字段如实让用户补。
5. 只在能确定/需要更新时给出该字段；拿不准的字段宁可不写。
6. question 自然简短，先补还没写好的关键小节；都齐时请用户确认并说“生成简历”。
7. 字段值用 \\n 分点分段，可直接排入简历对应小节。"""


def _docx_plain_text(path: Path, limit: int = 4000) -> str:
    """把 docx 的正文/表格文本粗略拼成纯文本（供 AI 当素材参考）。"""
    try:
        doc = DocxDocument(str(path))
        parts = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(p for p in parts if p and p.strip())[:limit]
    except Exception:
        return ""


def _load_resume_text(resume_id: str | None) -> str:
    if not resume_id:
        return ""
    try:
        from ..database import get_db
        with get_db() as db:
            row = db.execute("SELECT parsed_data FROM resumes WHERE id = ?", (resume_id,)).fetchone()
        return json.dumps(json.loads(row["parsed_data"] or "{}"), ensure_ascii=False)[:4000] if row else ""
    except Exception:
        return ""


async def _ai_extract(session: dict[str, Any], user_text: str) -> tuple[dict[str, str], str]:
    """调用 LLM 产出 (fields 更新, 下一句提问)。"""
    template_id = session["template_id"]
    fields = get_template(template_id).get("fields") or []
    current = session.get("fields") or {}
    history = session.get("history") or []

    lines = ["\n".join(f"- {i + 1}. {f}" for i, f in enumerate(fields))]
    prompt = (
        "需要填写的简历占位字段：\n" + "\n".join(f"- {f}" for f in fields) +
        "\n\n当前已填：\n" + json.dumps(current, ensure_ascii=False) +
        "\n\n对话记录（最近几条）：\n" + "\n".join(
            f"{h.get('role')}: {h.get('content')}" for h in history[-6:]
        ) +
        "\n\n用户刚刚说：\n" + user_text +
        "\n\n请按规则输出 JSON。"
    )
    resume_text = session.get("resume_text") or ""
    if resume_text:
        prompt += "\n\n已有简历素材（仅参考润色，不编造）：\n" + resume_text

    raw = await llm_generate(prompt=prompt, system=_SYSTEM, temperature=0.5)
    start, end = raw.find("{"), raw.rfind("}")
    if start == -1 or end <= start:
        raise ValueError("AI 返回格式异常，请再说一次")
    obj = json.loads(raw[start : end + 1])
    fields_update = {str(k).strip(): str(v).strip() for k, v in (obj.get("fields") or {}).items() if v}
    question = str(obj.get("question") or "").strip()
    return fields_update, question


def _missing(fields: list[str], current: dict[str, str]) -> list[str]:
    return [f for f in fields if not (current.get(f) or "").strip()]


def create_session(template_id: str, resume_id: str | None = None) -> dict[str, Any]:
    tpl = get_template(template_id)
    if not tpl:
        raise ValueError("模板不存在")
    session_id = generate_id()
    material = _load_resume_text(resume_id)
    if not material:
        # 未绑定“拜帖”简历时，用上传的这份（版式）简历原文作润色素材
        material = _docx_plain_text(TEMPLATE_DIR / f"{template_id}.docx")
    session = {
        "session_id": session_id,
        "template_id": template_id,
        "resume_id": resume_id,
        "fields": {},
        "history": [],
        "status": "collecting",
        "file_name": None,
        "created_at": now_iso(),
        "updated_at": now_iso(),
        "resume_text": material,
    }
    idx = _read_json(SESSION_INDEX)
    idx[session_id] = session
    _write_json(SESSION_INDEX, idx)
    return session


def get_session(session_id: str) -> dict[str, Any] | None:
    return _read_json(SESSION_INDEX).get(session_id)


def _save_session(session: dict[str, Any]) -> None:
    session["updated_at"] = now_iso()
    idx = _read_json(SESSION_INDEX)
    idx[session["session_id"]] = session
    _write_json(SESSION_INDEX, idx)


async def chat(session_id: str, user_text: str) -> dict[str, Any]:
    session = get_session(session_id)
    if not session:
        raise ValueError("会话不存在")
    fields = get_template(session["template_id"]).get("fields") or []

    fields_update, question = await _ai_extract(session, user_text)
    merged = dict(session.get("fields") or {})
    merged.update(fields_update)
    session["fields"] = merged
    session["history"] = (session.get("history") or []) + [
        {"role": "user", "content": user_text[:2000]},
        {"role": "assistant", "content": question[:2000]},
    ]
    session["history"] = session["history"][-12:]

    missing = _missing(fields, merged)
    done = not missing
    if done:
        session["status"] = "ready"
        reply = "所有字段已齐备，可以点击「生成简历」了；也可以继续补充我遗漏的细节。"
        question = ""
    else:
        reply = question or f"还缺：{'、'.join(missing)}，请继续告诉我。"
    _save_session(session)
    return {"session_id": session_id, "fields": merged, "missing": missing, "done": done, "question": reply}


def generate_docx(session_id: str) -> dict[str, Any]:
    session = get_session(session_id)
    if not session:
        raise ValueError("会话不存在")
    tpl = get_template(session["template_id"])
    if not tpl:
        raise ValueError("模板不存在")
    is_builtin = tpl.get("mode") == "builtin"
    if not is_builtin:
        src = TEMPLATE_DIR / f"{session['template_id']}.docx"
        if not src.exists():
            raise ValueError("模板文件缺失")
    out = GEN_DIR / f"{session_id}.docx"
    fields_values = session.get("fields") or {}
    key = (tpl or {}).get("builtin_key") or "classic"
    photo = session.get("photo_file")
    photo_path = str(GEN_DIR / photo) if photo else None
    # 统一用内置专业模板出稿（布局可控，输出整齐）；上传文件只作内容参考
    out.write_bytes(render_builtin_docx(key, fields_values, photo_path))
    name = fields_values.get("姓名") or ""
    # 清洗文件名中 Windows 不允许的字符
    safe_name = re.sub(r'[\\/:*?"<>|\r\n\t]+', "_", str(name)).strip().strip(".") or ""
    session["status"] = "ready"
    session["file_name"] = f"简历-{safe_name}.docx" if safe_name else "简历.docx"
    _save_session(session)
    return {"session_id": session_id, "path": str(out), "file_name": session["file_name"]}


def _desktop_dir() -> Path | None:
    """系统桌面路径（Windows 走注册表以兼容 OneDrive 重定向）。"""
    if not settings.resume_save_desktop:
        return None
    cand: Path | None = None
    if os.name == "nt":
        try:
            import winreg
            key = winreg.OpenKey(
                winreg.HKEY_CURRENT_USER,
                r"Software\Microsoft\Windows\CurrentVersion\Explorer\User Shell Folders",
            )
            try:
                val, _ = winreg.QueryValueEx(key, "Desktop")
            finally:
                winreg.CloseKey(key)
            cand = Path(os.path.expandvars(str(val)))
        except Exception:
            cand = None
    if cand is None:
        cand = Path.home() / "Desktop"
    return cand if cand.is_dir() else None


def save_to_desktop(src: Path, file_name: str) -> dict[str, Any] | None:
    """把生成的简历另存一份到系统桌面（尽力而为，返回桌面文件名/路径）。"""
    desktop = _desktop_dir()
    if desktop is None:
        return None
    stem = Path(file_name).stem or "简历"
    suffix = Path(file_name).suffix or ".docx"
    target = desktop / file_name
    n = 1
    while target.exists():
        target = desktop / f"{stem} ({n}){suffix}"
        n += 1
    try:
        shutil.copyfile(src, target)
    except Exception:
        logger.warning("保存简历到桌面失败：%s", target, exc_info=True)
        return None
    return {"filename": target.name, "path": str(target)}
